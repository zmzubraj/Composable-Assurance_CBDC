from __future__ import annotations
from pathlib import Path
import json, math, hashlib, textwrap
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
RES=ROOT/'results'; FIG=ROOT/'figures'; OUT=ROOT/'output'; OUT.mkdir(exist_ok=True)
DOCX=OUT/'Composable_Assurance_CBDC_National_Qualification_Manuscript_v7.docx'

# Evidence inputs
proto=json.loads((RES/'cross_border_bft_v5.json').read_text())
model=json.loads((RES/'cross_border_model_v5.json').read_text())
priv=json.loads((RES/'privacy_v7_experiment_summary.json').read_text())
dp=json.loads((RES/'dp_v7_summary.json').read_text())
aml=json.loads((RES/'aml_v7_summary.json').read_text())
san=json.loads((RES/'sanctions_v7_summary.json').read_text())
econ=json.loads((RES/'economic_v7_summary.json').read_text())
perf=json.loads((RES/'performance_v7_summary.json').read_text())
privacy_summary=pd.read_csv(RES/'privacy_v7_summary.csv')
privacy_rank=pd.read_csv(RES/'privacy_v7_ranking_summary.csv')
dp_summary=pd.read_csv(RES/'dp_v7_summary_table.csv')
aml_family=pd.read_csv(RES/'aml_v7_family_effects.csv')
aml_pairs=pd.read_csv(RES/'aml_v7_paired_effects.csv')
aml_work=pd.read_csv(RES/'aml_v7_workload.csv')
san_policy=pd.read_csv(RES/'sanctions_v7_policy_points.csv')
san_kind=pd.read_csv(RES/'sanctions_v7_by_kind.csv')
san_work=pd.read_csv(RES/'sanctions_v7_workload.csv')
econ_ceiling=pd.read_csv(RES/'economic_v7_identified_ceiling.csv')
econ_candidates=pd.read_csv(RES/'economic_v7_candidate_set.csv')
econ_flow=pd.read_csv(RES/'economic_v7_flow_duration.csv')
perf_scen=pd.read_csv(RES/'performance_v7_scenarios.csv')
perf_units=pd.read_csv(RES/'performance_v7_capacity_lower_bounds.csv')

# Theme
BLUE='1F4E79'; MID='5B9BD5'; LIGHT='EAF2F8'; PALE='F7F9FB'; ORANGE='FCE4D6'; GREEN='E2F0D9'; GRAY='5B6573'; WHITE='FFFFFF'; DARK='222222'; RED='9C0006'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def margin(cell,v=70):
    tcPr=cell._tc.get_or_add_tcPr(); m=tcPr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tcPr.append(m)
    for n in ['top','start','bottom','end']:
        x=m.find(qn('w:'+n))
        if x is None: x=OxmlElement('w:'+n); m.append(x)
        x.set(qn('w:w'),str(v)); x.set(qn('w:type'),'dxa')

def set_cell(cell,text,bold=False,size=7.65,color=DARK,align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    r=p.add_run(str(text)); r.bold=bold; r.font.name='Times New Roman'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margin(cell)

def repeat_header(row):
    x=OxmlElement('w:tblHeader'); x.set(qn('w:val'),'true'); row._tr.get_or_add_trPr().append(x)

def table(doc,heads,rows,caption=None,font=7.45,widths=None):
    if caption:
        p=doc.add_paragraph(caption,style='Caption'); p.paragraph_format.keep_with_next=True
    t=doc.add_table(rows=1,cols=len(heads)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    repeat_header(t.rows[0])
    for i,h in enumerate(heads):
        set_cell(t.rows[0].cells[i],h,True,font,WHITE,WD_ALIGN_PARAGRAPH.CENTER); shade(t.rows[0].cells[i],BLUE)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell(cells[i],v,False,font); shade(cells[i],WHITE if ri%2==0 else PALE)
    if widths:
        for row in t.rows:
            for c,w in zip(row.cells,widths): c.width=Inches(w)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
    return t

def fig(doc,path,caption,width=6.65):
    path=Path(path)
    if not path.exists(): return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True; p.paragraph_format.space_after=Pt(1)
    sh=p.add_run().add_picture(str(path),width=Inches(width)); sh._inline.docPr.set('title',path.stem.replace('_',' ')); sh._inline.docPr.set('descr',caption)
    c=doc.add_paragraph(caption,style='Caption'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER

def heading(doc,text,level=1):
    p=doc.add_heading(text,level=level); p.paragraph_format.keep_with_next=True; return p

def body(doc,text):
    return doc.add_paragraph(text,style='Body Text')

def bullets(doc,items):
    for x in items: doc.add_paragraph(x,style='List Bullet')

def equation(doc,text,n):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(f'{text}    ({n})'); r.font.name='Cambria Math'; r.font.size=Pt(9.7)

def callout(doc,title,text,fill=GREEN):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; repeat_header(t.rows[0]); shade(t.cell(0,0),fill); margin(t.cell(0,0),90)
    p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(title+' '); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def page_break(doc): doc.add_page_break()

def fmt(x,n=3): return f'{float(x):.{n}f}'

def make():
    d=Document(); s=d.sections[0]
    s.page_height=Cm(29.7); s.page_width=Cm(21); s.top_margin=Cm(1.55); s.bottom_margin=Cm(1.55); s.left_margin=Cm(1.65); s.right_margin=Cm(1.65); s.header_distance=Cm(.65); s.footer_distance=Cm(.65)
    st=d.styles
    st['Normal'].font.name='Times New Roman'; st['Normal'].font.size=Pt(9.35)
    st['Body Text'].font.name='Times New Roman'; st['Body Text'].font.size=Pt(9.35); st['Body Text'].paragraph_format.space_after=Pt(3.7); st['Body Text'].paragraph_format.line_spacing=1.03
    for name,size in [('Title',19),('Subtitle',10.5),('Heading 1',13.5),('Heading 2',11.2),('Heading 3',10.1)]:
        x=st[name]; x.font.name='Arial'; x.font.size=Pt(size); x.font.color.rgb=RGBColor.from_string(BLUE); x.font.bold=True; x.paragraph_format.space_before=Pt(6); x.paragraph_format.space_after=Pt(2.5); x.paragraph_format.keep_with_next=True
    st['Caption'].font.name='Arial'; st['Caption'].font.size=Pt(7.5); st['Caption'].font.italic=True; st['Caption'].font.color.rgb=RGBColor.from_string(GRAY); st['Caption'].paragraph_format.space_after=Pt(3)
    st['List Bullet'].font.name='Times New Roman'; st['List Bullet'].font.size=Pt(9.2); st['List Bullet'].paragraph_format.space_after=Pt(2)
    # running header/footer
    hp=s.header.paragraphs[0]; hp.text='Composable Assurance for Sovereign Digital Currency'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; hp.runs[0].font.size=Pt(7.1); hp.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
    fp=s.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run('Md Zubaer Mahmood Zubraj | 6 August 2026 | '); r.font.size=Pt(7.1); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

    # Article front matter, no decorative cover page
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(3)
    r=p.add_run('Composable Assurance for Sovereign Digital Currency'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(5)
    r=p.add_run('Privacy, financial integrity, adaptive policy and national-scale qualification across independent sovereign ledgers'); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(GRAY)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(1); p.add_run('Md Zubaer Mahmood Zubraj').bold=True
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(5); p.add_run('Independent Researcher | Corresponding author: zubraj14@gmail.com').font.size=Pt(8.6)

    heading(d,'Abstract')
    body(d,
         f'A production central bank digital currency (CBDC) must preserve monetary singleness, privacy, financial integrity, legal finality, bank intermediation and operational resilience as one composed system. This paper presents an evidence-constrained reference architecture for reserve-settled domestic conversion and payment-versus-payment across independently governed sovereign ledgers. Cross-border finalization uses bilateral compliance authorizations, final domestic prepare receipts and 5-of-7 PREPARE and terminal COMMIT-or-ABORT certificates; an abstract model explores {model["reachable_states"]:,} states and {model["transition_edges"]:,} transitions without an encoded conflicting-certificate or split-finality state under the declared two-key equivocation bound. Privacy is evaluated with learned network and compromised-intermediary attackers, unseen users, an independently generated behavioral regime and candidate sets up to 3,000. Under generator shift, network-observer AUC falls from 0.950 with pseudonym rotation alone to 0.605 with shielding and batching, but unseen-user AUC remains 0.830, demonstrating residual metadata leakage rather than perfect anonymity. Registered public statistics use person-level add/remove adjacency, contribution bound three, Laplace sensitivity three and a twelve-release privacy accountant; at annual epsilon four, dense cells have mean absolute error 9.46 while sparse cells retain 25.6% median relative error. AML/CFT evaluation uses sixteen independently generated graphs from four families. Cross-institution network features improve mean PR-AUC by 0.058 (bootstrap 95% interval 0.021-0.093; Wilcoxon p=0.013), but the effect is positive in 12 of 16 graphs and heterogeneous by family, so operational effectiveness remains a prospective outcome question. A sanctions benchmark indexes 19,178 official records and 20,010 aliases and evaluates calibrated high-recall, balanced and low-workload policies. The balanced policy obtains recall 0.874 and false-positive rate 0.00357; at 0.01% prevalence this implies only 2.39% positive predictive value and about 3,571 false alerts per million screenings, making analyst capacity and legal review part of system design. Economic analysis replaces a universal holding-limit claim with partial identification and robust policy selection; a 5% gross-exposure tolerance with 300 million fully participating users yields an illustrative ceiling of EUR 1,643, not a recommendation. Finally, a trace-driven tandem-queue digital twin satisfies a candidate two-second domestic p99 objective in declared 2,000-20,000 TPS scenarios but fails at 40,000 TPS with one region unavailable (p99 8.43 s). The capacity equation is treated as a necessary lower bound, while national-scale evidence requires physical multi-region, certified-HSM and independent pilot qualification. The result is a falsifiable research and deployment pathway rather than an unsupported production claim.')
    body(d,'Keywords: central bank digital currency; cross-border settlement; differential privacy; AML/CFT; sanctions screening; financial stability; queueing; operational resilience')

    heading(d,'1. Introduction')
    body(d,'CBDC work has advanced from conceptual ledgers to two-tier prototypes, privacy mechanisms, offline payment designs, programmable compliance and cross-border experiments. The 2024 BIS survey reported broad central-bank exploration, but exploration does not imply that one architecture has jointly established monetary correctness, privacy, financial-integrity effectiveness, economic stability and national-scale resilience [1]. The central challenge is compositional assurance: improving one property can weaken another. Detailed transaction visibility may assist investigation but create population-scale surveillance; aggressive privacy may weaken typology detection; immediate convertibility may increase usability while accelerating deposit flight; cross-system coordination may improve reach while creating a shared governance or failure concentration.')
    body(d,'This paper studies the following question: under which explicitly testable assumptions can independent sovereign CBDC systems provide reserve-settled conversion, privacy-preserving financial integrity, evidence-bound cross-border finality, adaptive policy limits and a falsifiable path to national-scale qualification? The intended contribution is not a new consensus algorithm or a universal monetary-policy number. It is a composition and evidence methodology that connects each architecture claim to a formal, empirical, operational or institutional test.')
    bullets(d,[
        'A sovereign-boundary architecture in which no foreign ledger or interlink can mint, burn or directly mutate another currency.',
        'A cross-border evidence protocol that binds payment intent, unequal-currency FX quote, bilateral compliance authorizations and final prepare receipts to explicit terminal certificates.',
        'A privacy evaluation with learned attackers, distribution shift and large ranking populations, plus a separately scoped person-level differential-privacy service for public statistics.',
        'Leakage-controlled AML/CFT experiments across independent graph families, calibrated sanctions-screening policy frontiers and prospective effectiveness outcomes.',
        'A partially identified, robust and adaptive policy-limit method that treats stability, utility, inclusion and distribution as constraints rather than assuming one transferable optimum.',
        'A four-stage national qualification method: analytical lower bounds, trace-driven queue simulation, physical multi-region testing and an independently observed pilot.'
    ])
    fig(d,FIG/'v7_assurance_stack.png','Figure 1. Composable assurance and the evidence ladder used throughout the paper.',6.7)

    heading(d,'2. Related work and novelty boundary')
    body(d,'Project Icebreaker demonstrated cross-border retail settlement through two domestic CBDC legs funded by FX providers [2]. Project Mandala developed rules and proof engines for compliance-by-design [3]. Project Agorá demonstrated atomic multi-currency wholesale settlement on a shared programmable platform [4]. Project FuSSE reported 10,000 transactions per second in a controlled modular proof of concept while explicitly identifying scaling and operational challenges [5]. These projects provide essential components, but they use different market structures and evidence scopes.')
    body(d,'The transaction-decision service builds on atomic-commit and fault-tolerant consensus literature rather than claiming a new consensus primitive [6]-[8]. Privacy design follows the institutional and technical separation emphasized by NIST, IMF, Tourbillon and W3C work [9], [22], [23], [33]. Financial-integrity evaluation incorporates the network perspectives explored by Aurora and Hertha [19], [20] while distinguishing predictive performance from the effectiveness outcomes required by FATF [10], [26], [27], [34]. Monetary-policy analysis is informed by BIS and ECB work on gradual and stress-driven bank disintermediation [14]-[16]. Operational qualification uses TIPS, Pix and FuSSE as workload references without treating any one reference as a universal CBDC demand forecast [5], [17], [18].')
    table(d,['Prior stream','Established contribution','Boundary addressed in this paper'],[
        ['Icebreaker [2]','Two domestic cross-border legs and FX-provider competition','Evidence-bound prepare and terminal certificates; explicit recovery and legal evidence'],
        ['Mandala [3]','Rules and cryptographic compliance proofs','Composition with monetary accounting, privacy, AML analytics and settlement finality'],
        ['Agorá [4]','Atomic multi-currency wholesale settlement on a shared platform','Independent-ledger isolation and jurisdiction-specific policy control'],
        ['FuSSE [5]','10,000 TPS controlled modular settlement proof of concept','Qualification method that separates analytical, simulated, physical and pilot evidence'],
        ['Paxos Commit / BFT [6]-[8]','Fault-tolerant decision and consensus foundations','CBDC evidence objects, sovereign boundaries and legal mapping'],
        ['BIS/ECB policy analysis [14]-[16]','Structural and bank-level CBDC exposure analysis','Partially identified, adaptive rule rather than universal transfer of one estimate']
    ],'Table 1. Principal novelty boundaries.',7.15)
    callout(d,'Scoped novelty.', 'The contribution is the explicit composition of monetary accounting, privacy, financial integrity, independent-ledger finality, adaptive policy and evidence-gated scale qualification. Individual primitives are not claimed as novel.',GREEN)

    heading(d,'3. System model, requirements and monetary invariants')
    fig(d,FIG/'v7_architecture.png','Figure 2. Technology-neutral architecture and the sovereign boundary.',6.7)
    body(d,'Each domestic system contains a central-bank monetary core, a reserve/RTGS adapter, regulated payment-interface providers (PIPs), identity and credential services, financial-integrity services, an offline-risk module and operational-security controls. The implementation may use a high-assurance replicated database or a permissioned deterministic ledger. Governance remains centralized in law and monetary authority even when technical operation is distributed for resilience.')
    table(d,['Requirement','Operational interpretation'],[
        ['Monetary singleness','CBDC is redeemable at par and recorded as a direct central-bank liability.'],
        ['Supply conservation','Issuance, redemption, escrow and offline allocation reconcile to unique sources and immutable receipts.'],
        ['Two-tier accountability','PIPs retain customer-facing identity, support, fraud and AML responsibilities.'],
        ['Purpose limitation','Each actor receives only the data required for its legal and operational role.'],
        ['Deterministic finality','Terminal monetary state is technically and legally recognized; correction uses a new compensating transaction.'],
        ['Failure containment','A foreign, PIP or analytics failure cannot create money or corrupt unrelated domestic balances.'],
        ['Evidence before scale','Scope expands only after formal, privacy, AML, sanctions, economic, performance and legal gates pass.']
    ],'Table 2. Core design requirements.',7.35)
    body(d,'Wholesale issuance converts reserve liabilities into CBDC inventory; retail conversion exchanges a commercial-bank deposit for existing CBDC inventory. These operations must not be conflated. Let R denote reserve liabilities, C issued CBDC, N notes, I_b the CBDC inventory of bank b, W_u a customer wallet and E settlement escrow.')
    equation(d,'Delta C = - Delta R','1')
    equation(d,'R + C + N = constant  (for a pure liability-composition change)','2')
    equation(d,'C = sum_b I_b + sum_u W_u + E + O + A','3')
    body(d,'Here O is risk-bounded offline value and A contains explicitly defined system accounts. Every increase in C references one final reserve or authorized accounting event. Replays are idempotent, balances cannot become negative, and finalized history is not rolled back.')

    heading(d,'4. Evidence-bound cross-border settlement')
    fig(d,FIG/'v7_cross_border_sequence.png','Figure 3. Cross-border sequence with explicit PREPARE and terminal decision certificates.',6.7)
    body(d,'Currency A remains on ledger A and currency B remains on ledger B. A licensed FX provider is prefunded in both systems. The payer transfers A domestically to the provider; the provider transfers B domestically to the payee. The interlink coordinates signed evidence but never custodies or represents either CBDC.')
    table(d,['Object','Mandatory binding'],[
        ['Payment intent PI','payer/payee routing, currencies, maximum debit, minimum credit, corridor, expiry and unique transaction ID'],
        ['FX quote Q','unequal amounts, rate, fee, provider, validity and signature'],
        ['Compliance CA / CB','jurisdiction, policy and list versions, payment digest, amount and expiry'],
        ['Prepare PA / PB','final domestic escrow state, exact currency/amount, source, destination and finality receipt'],
        ['Prepare certificate PC','5-of-7 decision-node votes over the identical evidence digest'],
        ['Terminal certificate TC','5-of-7 COMMIT or ABORT votes; honest nodes durably record one terminal decision']
    ],'Table 3. Cross-border protocol objects.',7.15)
    body(d,'For n=7 signers and q=5, any two quorums intersect in at least 2q-n=3 nodes. If at most f=2 keys equivocate and every honest signer durably signs at most one terminal decision for a transaction, conflicting valid terminal certificates require at least one honest signer to violate its lock.')
    equation(d,'|Q_1 intersection Q_2| >= 2q - n = 3 > f = 2','4')
    table(d,['Property','Argument and boundary'],[
        ['No local-timeout abort','A domestic ledger releases escrow only on a valid ABORT certificate; silence or local time is insufficient.'],
        ['No committed-versus-aborted finality','Both ledgers verify a terminal certificate bound to the same digest before finalization.'],
        ['Compliance atomicity','PI, Q, CA, CB, PA and PB are included in the certificate digest.'],
        ['Replay resistance','Unique IDs, expiry, digest binding and idempotent state transitions reject duplicate economic effects.'],
        ['Sovereign isolation','The interlink cannot mint, burn, freeze or directly transfer either domestic currency.'],
        ['Proof boundary','Safety assumes signature security, durable honest locks, correct verification and no more than two equivocating keys; indefinite denial of service and dynamic membership are outside the model.']
    ],'Table 4. Cross-border properties and proof boundary.',7.05)
    body(d,f'The executable abstract model reaches {model["reachable_states"]:,} states through {model["transition_edges"]:,} transitions and finds no encoded conflicting-certificate or split-finality state. The eleven-service laboratory prototype completes {proto["completed_cross_border_transfers"]} unequal-currency transfers and rejects stale compliance, mismatched quotes and conflicting certificates. These results establish bounded model and laboratory evidence, not certified multi-region or Byzantine deployment.')

    page_break(d)
    heading(d,'5. Privacy architecture and registered differential privacy')
    fig(d,FIG/'v7_privacy_dp_pipeline.png','Figure 4. Institutional separation, learned privacy attacks and registered public-statistics releases.',6.7)
    heading(d,'5.1 Formal objective and attacker model',2)
    body(d,'The privacy objective is data minimization and bounded linkability, not universal anonymity. Full identity remains with the responsible PIP. The core ledger receives rotating references and minimum settlement state. Compliance presentations are bound to a transaction and must not expose stable presentation identifiers. Lawful re-identification requires a documented legal basis, an authorization quorum and an immutable access record.')
    table(d,['Attacker','Observable information','Primary inference risk'],[
        ['Ledger operator','rotating references, amount, time, state and counterparty class','behavioral linkage across rotations'],
        ['Network observer','endpoint class, transport timing, cadence and size bands','traffic correlation'],
        ['Merchant coalition','repeat merchant visits, time and spending pattern','cross-merchant profiling'],
        ['Compromised PIP','identity mapping and local history plus auxiliary data','direct re-identification and cross-system linkage']
    ],'Table 5. Privacy attacker model.',7.35)
    body(d,'Let b be a hidden bit selecting whether two observations originate from the same person. An attacker A receives only the declared leakage L and returns a guess. The linkability advantage is:')
    equation(d,'Adv_link(A) = | Pr[A(L_0,L_1)=b] - 1/2 |','5')
    body(d,'The formal objective requires negligible advantage for protected contexts under the chosen cryptographic and traffic assumptions. The empirical experiment deliberately relaxes those ideal assumptions and measures how much metadata remains exploitable.')

    heading(d,'5.2 Learned red-team evaluation',2)
    fig(d,FIG/'privacy_v7_learned_attack.png','Figure 5. Learned pairwise-attack ROC-AUC under unseen-user and independently generated evaluation regimes.',6.45)
    rows=[]
    for profile in ['rotation_only','relay_standardized','shielded_batched']:
        for split in ['independent_generator','unseen_users']:
            r=privacy_summary[(privacy_summary.profile==profile)&(privacy_summary.attacker=='network')&(privacy_summary.split==split)].iloc[0]
            rows.append([profile.replace('_',' '),split.replace('_',' '),f'{r.auc_mean:.3f}',f'{r.auc_min:.3f}-{r.auc_max:.3f}'])
    table(d,['Privacy profile','Evaluation split','Network-observer AUC','Seed range'],rows,'Table 6. Learned network-observer linkage results.',7.3)
    body(d,'The attack uses traffic timing, endpoint classes, merchant repetition, wallet implementation fingerprints and behavioral summaries. The independent-generator split changes the calendar/payroll regime, while the unseen-user split holds out identities but preserves the generator. Pseudonym rotation alone is highly linkable. Relay standardization reduces linkability, and shielding plus batching reduces it further, but unseen-user AUC remains materially above random. The experiment therefore rejects any claim that pseudonym rotation or batching alone establishes anonymity.')
    rankrows=[]
    for profile in ['rotation_only','relay_standardized','shielded_batched']:
        r=privacy_rank[(privacy_rank.profile==profile)&(privacy_rank.split=='independent_generator')&(privacy_rank.candidate_size==3000)].iloc[0]
        rankrows.append([profile.replace('_',' '),f'{r.top1_mean:.3f}',f'{r.top10_mean:.3f}',f'{r.mrr_mean:.3f}',f'{r.median_rank_mean:.1f}'])
    table(d,['Profile','Top-1 at 3,000','Top-10','MRR','Median rank'],rankrows,'Table 7. Ranking attack under independently generated traffic.',7.25)

    heading(d,'5.3 Registered public-statistics differential privacy',2)
    body(d,'Differential privacy is used only for registered aggregate publications, not as a blanket description of the ledger or AML system. Two datasets are adjacent when one person and all of that person’s registered contributions are added or removed. A person contributes one count to at most three cells per monthly release, so the L1 sensitivity of the histogram is three. The release mechanism is:')
    equation(d,'M(D) = max(0, f(D) + Laplace(Delta_1 / epsilon_t)),   Delta_1 = 3','6')
    equation(d,'epsilon_total = sum_(t=1)^12 epsilon_t = 4,   epsilon_t = 1/3','7')
    body(d,'Non-negative clipping is post-processing and does not weaken the formal guarantee, but it creates bias in sparse cells and is therefore measured. Ten thousand neighboring-dataset checks confirm a maximum L1 change of three in the implementation.')
    fig(d,FIG/'dp_v7_utility.png','Figure 6. Utility and clipping effects for dense and sparse registered histograms.',6.35)
    table(d,['Annual epsilon','Dense MAE','Dense median relative error','Sparse MAE','Sparse median relative error','Sparse clipping bias'],[
        [int(e),f'{dp_summary[(dp_summary.epsilon_total==e)&(dp_summary.regime=="dense")].iloc[0].mae:.2f}',f'{100*dp_summary[(dp_summary.epsilon_total==e)&(dp_summary.regime=="dense")].iloc[0].median_relative_error:.1f}%',f'{dp_summary[(dp_summary.epsilon_total==e)&(dp_summary.regime=="sparse")].iloc[0].mae:.2f}',f'{100*dp_summary[(dp_summary.epsilon_total==e)&(dp_summary.regime=="sparse")].iloc[0].median_relative_error:.1f}%',f'{dp_summary[(dp_summary.epsilon_total==e)&(dp_summary.regime=="sparse")].iloc[0].clipping_bias:.2f}'] for e in [1,2,4,8]
    ],'Table 8. Differential-privacy utility by annual budget.',7.1)
    callout(d,'Guarantee boundary.', 'The guarantee covers registered public histograms and their accountant. It does not cover raw data inside the trusted curator, identity maps, AML graphs, case files, arbitrary dashboards or unregistered revisions.',ORANGE)

    heading(d,'6. AML/CFT evaluation: predictive evidence and effectiveness gate')
    fig(d,FIG/'v7_aml_evaluation.png','Figure 7. Leakage-controlled AML comparison and prospective effectiveness outcomes.',6.65)
    body(d,'FATF evaluates effectiveness separately from technical compliance [10]. Accordingly, the experiment asks a narrower predictive question: do minimized cross-institution network features improve account-level ranking relative to a genuinely PIP-local view across independent graph-generating regimes? It does not infer that a higher PR-AUC automatically improves investigations, suspicious-transaction reports or asset recovery.')
    table(d,['Design element','Implementation'],[
        ['Independent experimental unit','Sixteen independently generated graphs; four graphs in each of four generator families'],
        ['Families','profile mix, calendar/payroll, merchant network and remittance corridor'],
        ['PIP-local baseline','features use only transactions visible to the institution'],
        ['Network model','minimized cross-PIP graph features recomputed separately inside every train/test subgraph'],
        ['Leakage control','account-group split and graph-local feature construction; labels and typology membership do not cross partitions'],
        ['Robustness','family heterogeneity, adaptive dispersion/timing changes and workload-at-fixed-alert-budget'],
        ['Inference','graph-level paired differences, bootstrap interval, Wilcoxon test and exact sign test']
    ],'Table 9. AML/CFT predictive-evaluation design.',7.1)
    equation(d,'delta_i = AP_network,i - AP_local,i,   i = 1,...,16','8')
    fig(d,FIG/'aml_v7_families.png','Figure 8. Mean PR-AUC difference by independent generator family.',6.3)
    table(d,['Family','PIP-local mean AP','Network mean AP','Mean difference','SD of difference'],[
        [r.family.replace('_',' '),f'{r.local_mean:.3f}',f'{r.network_mean:.3f}',f'{r.delta_mean:+.3f}',f'{r.delta_sd:.3f}'] for _,r in aml_family.iterrows()
    ],'Table 10. Family-specific AML effects.',7.2)
    body(d,f'Across all sixteen graphs, the mean paired difference is +{aml["mean_pr_auc_delta"]:.3f}; the bootstrap 95% interval is {aml["bootstrap_95_ci"][0]:.3f} to {aml["bootstrap_95_ci"][1]:.3f}, the paired Wilcoxon p-value is {aml["wilcoxon_two_sided_p"]:.4f}, and the exact sign-test p-value is {aml["sign_test_two_sided_p"]:.4f}. Twelve of sixteen graphs show a positive effect. The merchant-network family has a near-zero mean and high dispersion, while remittance-corridor graphs show the largest average gain. This heterogeneity is a central result: network analytics can help, but not uniformly.')
    table(d,['Alert budget','Network recall','Network precision','PIP-local recall','PIP-local precision','Alerts / million'],[
        [f'{100*b:.1f}%',f'{aml_work[(aml_work.alert_budget==b)&(aml_work.model=="network")].iloc[0].mean_recall:.3f}',f'{aml_work[(aml_work.alert_budget==b)&(aml_work.model=="network")].iloc[0].mean_precision:.3f}',f'{aml_work[(aml_work.alert_budget==b)&(aml_work.model=="pip_local")].iloc[0].mean_recall:.3f}',f'{aml_work[(aml_work.alert_budget==b)&(aml_work.model=="pip_local")].iloc[0].mean_precision:.3f}',f'{int(aml_work[(aml_work.alert_budget==b)&(aml_work.model=="network")].iloc[0].alerts_per_million):,}'] for b in [0.005,0.01,0.02]
    ],'Table 11. Predictive workload at fixed alert budgets.',7.05)
    body(d,'Even where ranking improves, recall at small alert budgets remains low. Production evaluation must therefore be prospective and outcome-based: evidence sufficiency, investigator time, escalation, repeat alerts, SAR/STR quality, FIU feedback, restraint or recovery, network disruption, complaints and subgroup harm. Model outputs remain decision support and never automatic guilt determinations.')

    page_break(d)
    heading(d,'7. Sanctions evaluation as an operational decision system')
    fig(d,FIG/'v7_sanctions_workflow.png','Figure 9. Sanctions screening from official-list ingestion to analyst and ownership decisions.',6.7)
    body(d,'Sanctions screening is a staged decision process, not a string-similarity threshold. Official records and aliases are versioned; candidate retrieval uses multilingual, token and phonetic representations; an evidence model incorporates names, dates and places of birth, nationality, address and entity type; thresholds are selected on a calibration set; analysts decide disposition under the applicable programme and licensing rules; and ownership is evaluated separately, including direct and indirect aggregate control [11]-[13].')
    table(d,['Benchmark property','Implementation'],[
        ['Official input','19,178 OFAC records and 20,010 official aliases'],
        ['Positive split','sanctioned entities disjoint across training, calibration and test'],
        ['Perturbations','exact, official alias, OCR, token order, transliteration and sparse-attribute cases'],
        ['Negative population','generated multilingual legitimate identities and hard near matches'],
        ['Calibration','threshold selected before held-out test; Brier score and calibration error reported'],
        ['Ownership','500 randomized monotonicity property tests for direct and indirect aggregate ownership'],
        ['Boundary','no live customer population, independent multilingual analyst adjudication or live corporate registry']
    ],'Table 12. Sanctions evaluation design.',7.15)
    fig(d,FIG/'sanctions_v7_policy_frontier.png','Figure 10. Calibrated sanctions policy frontier: recall and false-positive trade-off.',6.3)
    table(d,['Policy','Threshold','Recall','False-positive rate','Benchmark precision'],[
        [r.policy.replace('_',' '),f'{r.threshold:.3f}',f'{r.recall:.3f}',f'{r.fpr:.5f}',f'{r.precision_benchmark:.3f}'] for _,r in san_policy.iterrows()
    ],'Table 13. Calibrated screening policies.',7.25)
    equation(d,'PPV(pi) = Recall*pi / [Recall*pi + FPR*(1-pi)]','9')
    equation(d,'Analyst hours = (true alerts + false alerts) * minutes_per_alert / 60','10')
    fig(d,FIG/'sanctions_v7_workload.png','Figure 11. Prevalence-adjusted analyst workload across calibrated policies.',6.3)
    workload_rows=[]
    for policy in ['high_recall','balanced','low_workload']:
        r=san_work[(san_work.policy==policy)&(san_work.prevalence==0.0001)].iloc[0]
        workload_rows.append([policy.replace('_',' '),f'{100*r.prevalence_adjusted_ppv:.2f}%',f'{r.true_alerts_per_million:.1f}',f'{r.false_alerts_per_million:,.0f}',f'{r.analyst_hours_per_million_at_7min:,.0f}'])
    table(d,['Policy at 0.01% prevalence','Operational PPV','True alerts / million','False alerts / million','Analyst hours / million'],workload_rows,'Table 14. Prevalence-adjusted workload illustration.',7.0)
    kind_rows=[]
    for _,r in san_kind.iterrows():
        kind_rows.append([str(r['kind']).replace('_',' '),f'{r["end_to_end_recall"]:.3f}',f'{r["n"]:.0f}'])
    table(d,['Held-out perturbation class','Recall at balanced policy','Positive queries'],kind_rows,'Table 15. Recall by evidence and perturbation class.',7.3)
    body(d,'The balanced policy obtains recall 0.874 and false-positive rate 0.00357 in the held-out benchmark. Benchmark precision is high because positives are deliberately overrepresented. At a 0.01% operational prevalence, the same rates imply only 2.39% positive predictive value and approximately 3,571 false alerts per million screenings. The low-workload policy reduces false alerts to approximately 714 per million but recall falls to 0.757. Sparse-attribute recall is 0.617. The real-world solution is therefore a risk-tiered frontier: high-recall policies for narrow high-risk cohorts, lower-workload policies for broad low-prevalence populations, analyst reason codes, workload caps, periodic calibration, and mandatory escalation when ownership or evidence is ambiguous.')

    heading(d,'8. Economic policy identification and adaptive guardrails')
    fig(d,FIG/'v7_economic_decision.png','Figure 12. From jurisdictional data to a partially identified and adaptively governed policy set.',6.65)
    body(d,'A universally correct CBDC holding limit cannot be inferred from public cross-country evidence. The relevant decision combines a holding threshold H, a system and user conversion-flow limit F, tiered remuneration r, automatic sweeps and liquidity facilities. The policy problem is partially identified because household demand, bank replacement funding, crisis behavior and distributional preferences vary across jurisdictions.')
    body(d,'A transparent first constraint bounds gross CBDC exposure. Let D be eligible household deposits, N the eligible population, a the adoption share and tau the maximum accepted gross exposure share. Then:')
    equation(d,'H <= tau D / (N a)','11')
    body(d,'Using the public euro-area deposit anchor D=EUR 9.861 trillion, N=300 million, a=1 and tau=5% yields an illustrative ceiling H<=EUR 1,643. This is not a policy recommendation: lowering adoption to 50% doubles the arithmetic ceiling, and selecting 3%, 7.5% or 10% exposure tolerances changes it materially. The purpose is to expose the normative and empirical assumptions instead of hiding them in a single number.')
    fig(d,FIG/'economic_v7_identified_set.png','Figure 13. Illustrative holding-limit ceilings across exposure tolerances, eligible populations and adoption shares.',6.35)
    table(d,['Holding limit','Gross exposure share (300m, full adoption)','Passes 5% cap','Passes 10% cap','Relation to BIS structural point'],[
        [f'EUR {int(r.holding_limit_eur):,}',f'{100*r.gross_exposure_share_300m_full:.2f}%',str(bool(r.passes_5pct_exposure_cap)),str(bool(r.passes_10pct_exposure_cap)),f'{r.distance_from_bis_point_eur:+.0f} EUR distance'] for _,r in econ_candidates[econ_candidates.holding_limit_eur.isin([500,1000,1500,1750,2000,2500,3000])].iterrows()
    ],'Table 16. Illustrative candidate limits under one exposure definition.',7.05)
    body(d,'The policy authority should choose within the feasible set by minimizing expected and tail loss over an uncertainty set U. Let theta=(H,F,r) and L combine bank funding stress, payment-utility loss, exclusion and distributional burden:')
    equation(d,'theta* = arg min_(theta in Theta) max_(P in U) { E_P[L(theta,S)] + lambda CVaR_alpha[L(theta,S)] }','12')
    body(d,'The optimization is subject to bank-level liquidity and funding constraints, ordinary-payment utility, accessibility, consumer-protection and legal constraints. A robust-regret or CVaR formulation does not create missing evidence; it makes uncertainty explicit. The parameter set must be updated on a predeclared schedule using survey data, pilot holdings, bank balance sheets, liquidity-facility use, deposit repricing, subgroup outcomes and stress simulations.')
    equation(d,'Minimum binding days = ceil(desired conversion share / daily system flow share)','13')
    table(d,['Desired conversion share','0.25% daily limit','0.5% daily limit','1% daily limit','2% daily limit'],[
        [f'{100*s:.0f}%',*[str(int(econ_flow[(econ_flow.desired_conversion_share==s)&(econ_flow.daily_system_flow_limit_share==f)].iloc[0].minimum_days_if_binding)) for f in [0.0025,0.005,0.01,0.02]]] for s in [0.02,0.05,0.09]
    ],'Table 17. Mechanical duration imposed by system-wide flow limits (days).',7.25)
    callout(d,'Policy conclusion.', 'The valid scientific output is a jurisdiction-specific feasible set, uncertainty interval and transparent update rule - not one global holding-limit number.',GREEN)

    page_break(d)
    heading(d,'9. National-scale qualification: from lower bounds to physical evidence')
    fig(d,FIG/'v7_national_qualification.png','Figure 14. Four-stage national-scale qualification and the claim boundary.',6.7)
    body(d,'A capacity equation cannot guarantee sufficient hardware because service times, queueing, replication, synchronization, HSM saturation, database contention, network loss and recovery interact nonlinearly. The paper therefore uses a hierarchy. Stage 1 derives a necessary analytical lower bound. Stage 2 runs a trace-driven tandem-queue digital twin with bursts and faults. Stage 3 executes the same predeclared workload on physical multi-region infrastructure with certified cryptographic hardware. Stage 4 verifies operations under a controlled external pilot. National-scale performance is demonstrated only after Stages 3 and 4 pass.')
    body(d,'For stage j with qualification arrival rate lambda, mean service demand s_j, headroom multiplier h and maximum planned utilization rho_max, a necessary lower bound on parallel units is:')
    equation(d,'n_j >= ceil(lambda h s_j / rho_max)','14')
    body(d,'This lower bound is intentionally not a guarantee. In the digital twin, each stage is represented as a finite-server queue with time-varying arrivals. Approximate stage utilization is:')
    equation(d,'rho_j(t) = lambda(t) s_j / n_j','15')
    body(d,'The system passes only when the empirical latency and recovery distributions satisfy predeclared service objectives under normal, burst, region-loss, HSM-slowdown and overload-control scenarios. The model includes API/authentication, policy/sanctions, HSM verification, monetary transition and durable audit stages. Its service demands are declared scenario inputs, not measurements of a production system.')
    fig(d,FIG/'performance_v7_scenarios.png','Figure 15. Trace-driven digital-twin p99 latency across workload and failure scenarios.',6.35)
    table(d,['Scenario','Observed TPS','Domestic p99','Cross-border p99','Region loss','Domestic SLO'],[
        [r.scenario,f'{r.observed_arrival_tps:,.0f}',f'{r.domestic_p99_ms:,.1f} ms',f'{r.cross_border_p99_ms:,.1f} ms','yes' if r.region_loss else 'no','pass' if r.domestic_slo_pass else 'FAIL'] for _,r in perf_scen.iterrows()
    ],'Table 18. Digital-twin scenario outcomes.',7.0)
    body(d,'The declared 2,000-20,000 TPS scenarios satisfy the candidate two-second domestic p99 objective in the model. A 40,000 TPS overload with one region unavailable reaches p99 8.43 seconds and fails. This failure is valuable: it demonstrates that the model can falsify a capacity claim instead of being tuned to pass every scenario. Cross-border p99 remains below 0.5 seconds in the model because its modeled certificate path is not the congested bottleneck; physical qualification must verify whether this remains true with production consensus and cryptography.')
    fig(d,FIG/'performance_v7_capacity_bounds.png','Figure 16. Analytical lower-bound units by stage and qualification workload.',6.25)
    table(d,['Qualification load','API/auth units','Policy/sanctions units','HSM units','Monetary units','Audit units'],[
        [f'{t:,} TPS',*[str(int(perf_units[(perf_units.qualification_tps==t)&(perf_units.stage==stage)].iloc[0].minimum_parallel_units)) for stage in ['API and authentication','Policy and sanctions fast path','HSM signature verification','Monetary-state transition','Durable audit append']]] for t in [2000,10000,20000]
    ],'Table 19. Necessary parallel-unit lower bounds (2x headroom, 60% planned utilization).',7.05)
    table(d,['Physical qualification dimension','Required evidence before a national claim'],[
        ['Topology','At least three independent regions/failure domains; measured inter-region distributions and operator map'],
        ['Cryptography','Certified HSM throughput, threshold ceremonies, queue saturation, rotation, revocation and recovery'],
        ['Workload','Open-loop and closed-loop PIP mix; realistic amount, identity, sanctions and reconciliation paths; sustained 2x observed peak'],
        ['Faults','Region loss, partitions, stale rules, disk corruption, signer/PIP outages, network degradation and controlled overload'],
        ['State safety','No monetary-invariant violation, duplicate economic effect, conflicting terminal state or unbounded escrow'],
        ['Operations','RTO/RPO, reconciliation, fraud/support operations, change control, incident exercises and audit closure'],
        ['External validation','Independent observer, reproducible load traces, public result summary and controlled pilot exit decision']
    ],'Table 20. Mandatory physical and pilot acceptance evidence.',7.05)
    callout(d,'National-scale status.', 'The manuscript supplies an analytical bound, a reproducible digital-twin falsification test and a complete physical qualification protocol. It does not claim that national-scale performance has already been demonstrated.',ORANGE)

    heading(d,'10. Security, legal finality and governance')
    body(d,'CBDC is critical national infrastructure. Security controls include threshold HSM keys, separation of duties, deterministic replication, immutable evidence, independent reconciliation, software bills of materials, reproducible builds, certificate transparency, privacy access logs and cryptographic agility [21], [24], [25], [31]. General-purpose smart contracts do not receive issuance authority or unrestricted balance-write access. User-authorized conditional payments use audited templates and escrow primitives.')
    table(d,['Risk','Control and evidence'],[
        ['Issuance or decision-key compromise','threshold HSMs, amount/corridor limits, offline roots, witnessed ceremonies and rapid revocation'],
        ['Ledger corruption or equivocation','replication, state commitments, independent read replicas, reconciliation and restoration exercises'],
        ['PIP compromise','signed customer intent, identity-vault segregation, least privilege and certificate revocation'],
        ['Privacy insider abuse','purpose-based access, threshold disclosure, immutable logs and independent privacy authority'],
        ['Model manipulation','data lineage, drift and evasion tests, reason codes, analyst review and model-risk governance'],
        ['Supply-chain attack','SBOM, signed artifacts, reproducible build, independent review and staged release'],
        ['Cryptographic transition','algorithm inventory, hybrid migration, re-signing plan and measured post-quantum overhead']
    ],'Table 21. Security controls and evidence.',7.15)
    body(d,'Technical finality is insufficient without legal recognition. A corridor agreement must define the evidentiary status of PI, Q, CA/CB, PA/PB, PC and TC; quorum governance; liability for wrongful issuance or refusal; insolvency and FX-provider default; data transfer and sanctions duties; emergency suspension; dispute forum; and the domestic legal effect of finalization. A common protocol does not require identical laws, but each jurisdiction must map shared evidence to its own legal finality and supervisory framework [28]-[30], [32].')

    heading(d,'11. Claims-to-evidence register, limitations and research agenda')
    table(d,['Claim','Evidence reported','Permitted conclusion','External gate'],[
        ['Cross-border safety','state exploration, quorum argument, negative tests and lab prototype','bounded safety under declared assumptions','formal refinement, certified HSM and physical multi-region fault testing'],
        ['Privacy','learned synthetic attacks, unseen users, generator shift and large ranking sets','measured residual metadata linkability','governed real traces and independent red team'],
        ['Differential privacy','explicit adjacency, sensitivity, utility and code-enforced accountant','formal guarantee for registered histograms','independent implementation and release-governance audit'],
        ['AML/CFT','16 independent graphs, four families and graph-level inference','heterogeneous predictive improvement','prospective investigator, FIU, supervisory, disruption and harm outcomes'],
        ['Sanctions','full official-list index, entity-disjoint test, calibration and prevalence workload','laboratory screening frontier','real customer distribution, multilingual adjudication and live ownership data'],
        ['Economic policy','public anchors, partial identification and robust decision rule','conditional feasible set and update method','jurisdictional microdata, structural estimation and pilot evidence'],
        ['National qualification','lower bound, digital twin, overload failure and acceptance protocol','falsifiable qualification methodology','physical regions, certified HSMs, sustained load and independent pilot']
    ],'Table 22. Claims-to-evidence register.',6.9)
    body(d,'The principal limitations are deliberate and material. Privacy and AML experiments remain synthetic. The sanctions negative population is generated and does not reproduce every language, transliteration or customer-data quality problem. The economic illustration does not estimate behavioral parameters from microdata. The performance model uses declared service demands and cannot capture every implementation interaction. The cross-border fault model excludes indefinite denial of service, dynamic membership and post-compromise recovery. These limitations constrain the claims but make the research falsifiable.')
    body(d,'The next empirical programme is therefore preregistered around six institutional studies: governed multi-PIP privacy red-teaming; prospective AML investigator and FIU outcomes; multilingual sanctions adjudication with real customer distributions; bank- and household-level economic estimation; physical multi-region/HSM performance qualification; and independent legal, cryptographic and operational audit.')

    heading(d,'12. Conclusion')
    body(d,'A credible CBDC publication must distinguish architecture from evidence and evidence from deployment. The proposed system preserves monetary accounting and sovereign ledgers, binds cross-border finality to explicit evidence, quantifies rather than assumes privacy, restricts differential privacy to a formally specified release service, evaluates AML and sanctions under realistic workload boundaries, and treats the economic policy limit as partially identified and adaptive. For scale, the manuscript replaces a misleading hardware guarantee with a necessary lower bound, a trace-driven falsification model, physical acceptance criteria and an independently observed pilot gate. This combination provides a clean, reproducible and real-world research pathway while refusing to label laboratory results as national certification.')

    heading(d,'References')
    refs=[
        '[1] A. Illes, A. Kosse and P. Wierts, “Advancing in tandem: results of the 2024 BIS survey on central bank digital currencies and crypto,” BIS Papers No. 159, Bank for International Settlements, 22 Aug. 2025.',
        '[2] BIS Innovation Hub, “Project Icebreaker: Breaking new paths in cross-border retail CBDC payments,” Mar. 2023.',
        '[3] BIS Innovation Hub, “Project Mandala: shaping the future of cross-border payments compliance,” Oct. 2024.',
        '[4] BIS Innovation Hub and Institute of International Finance, “Project Agorá: a shared programmable platform for tokenised wholesale cross-border payments,” 27 May 2026.',
        '[5] BIS Innovation Hub, “Project FuSSE: exploring flexible, scalable and secure settlement engines,” 29 Jan. 2026.',
        '[6] J. Gray and L. Lamport, “Consensus on transaction commit,” ACM Transactions on Database Systems, vol. 31, no. 1, pp. 133-160, 2006, doi:10.1145/1132863.1132867.',
        '[7] M. Castro and B. Liskov, “Practical Byzantine fault tolerance,” Proc. OSDI, pp. 173-186, 1999.',
        '[8] M. Yin, D. Malkhi, M. K. Reiter, G. G. Gueta and I. Abraham, “HotStuff: BFT consensus with linearity and responsiveness,” Proc. ACM PODC, pp. 347-356, 2019, doi:10.1145/3293611.3331591.',
        '[9] NIST, “Guidelines for Evaluating Differential Privacy Guarantees,” Special Publication 800-226, Mar. 2025.',
        '[10] Financial Action Task Force, “Methodology for Assessing Technical Compliance with the FATF Recommendations and the Effectiveness of AML/CFT/CPF Systems,” updated June 2026.',
        '[11] U.S. Treasury, Office of Foreign Assets Control, “Frequently Asked Questions on Advanced Sanctions List Standard,” accessed 6 Aug. 2026.',
        '[12] U.S. Treasury, Office of Foreign Assets Control, “How Sanctions List Search works,” FAQs 246-250, accessed 6 Aug. 2026.',
        '[13] U.S. Treasury, Office of Foreign Assets Control, “Entities Owned by Blocked Persons: 50 Percent Rule,” FAQs 398-402, accessed 6 Aug. 2026.',
        '[14] R. Bidder, T. Jackson and M. Rottner, “CBDC and banks: disintermediating fast and slow,” BIS Working Papers No. 1280, 8 Jul. 2025.',
        '[15] European Central Bank, “Technical data on the financial stability impact of the digital euro,” seminar technical data, 22 Oct. 2025.',
        '[16] European Central Bank, “Preparation phase of a digital euro: closing report,” 30 Oct. 2025.',
        '[17] European Central Bank, “TIPS Capacity Elements,” requirement TIPS.UR.10.120, Feb. 2024; capacity reconfirmed Feb. 2026.',
        '[18] Banco Central do Brasil, “Instant Payments System (SPI) Annual Report 2024,” 2025.',
        '[19] BIS Innovation Hub, “Project Aurora: the power of data, technology and collaboration to combat money laundering,” 2023.',
        '[20] BIS Innovation Hub, “Project Hertha: identifying financial crime patterns in payment systems,” 2025.',
        '[21] BIS Innovation Hub, “Project Polaris: a security and resilience framework for CBDC systems,” 2023.',
        '[22] BIS Innovation Hub, “Project Tourbillon: exploring privacy, scalability and quantum-safe cryptography for CBDC,” 2023.',
        '[23] W3C, “Data Integrity BBS Cryptosuites v1.0,” Candidate Recommendation Draft, accessed 6 Aug. 2026.',
        '[24] NIST, “Module-Lattice-Based Key-Encapsulation Mechanism Standard,” FIPS 203, 2024.',
        '[25] NIST, “Module-Lattice-Based Digital Signature Standard,” FIPS 204, 2024.',
        '[26] Financial Action Task Force, “Guidance on Digital Identity,” 6 Mar. 2020.',
        '[27] Financial Action Task Force, “Guidance on Beneficial Ownership of Legal Persons,” Mar. 2023.',
        '[28] CPMI, “Fast payment system interlinking,” updated report, Feb. 2026.',
        '[29] ISO, “ISO 20022: Universal financial industry message scheme,” current edition.',
        '[30] CPMI-IOSCO, “Principles for Financial Market Infrastructures,” Apr. 2012.',
        '[31] Basel Committee on Banking Supervision, “Principles for Operational Resilience,” Mar. 2021.',
        '[32] European Union, General Data Protection Regulation, Regulation (EU) 2016/679.',
        '[33] International Monetary Fund, “Central Bank Digital Currency Data Use and Privacy Protection,” Fintech Note, 2024.',
        '[34] International Monetary Fund, “Central Bank Digital Currencies and Financial Integrity,” Fintech Note, 2025.'
    ]
    for x in refs: body(d,x)


    d.core_properties.title='Composable Assurance for Sovereign Digital Currency'
    d.core_properties.subject='Privacy, financial integrity, adaptive policy and national-scale qualification across independent sovereign ledgers'
    d.core_properties.author='Md Zubaer Mahmood Zubraj'
    d.core_properties.keywords='CBDC, cross-border settlement, privacy, differential privacy, AML, sanctions, economic policy, national-scale qualification'
    d.save(DOCX); print(DOCX)

if __name__=='__main__': make()
