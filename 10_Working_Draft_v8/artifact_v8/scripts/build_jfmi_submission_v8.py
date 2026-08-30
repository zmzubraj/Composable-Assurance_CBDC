from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "Composable_Assurance_CBDC_Author_Draft_v8.docx"
OUT = ROOT / "output" / "jfmi_provisional"
MAIN = OUT / "Composable_Assurance_CBDC_JFMI_Blinded_Manuscript_v8.docx"
TITLE_PAGE = OUT / "Composable_Assurance_CBDC_JFMI_Title_Page_TEMPLATE_v8.docx"
LEGENDS = OUT / "Composable_Assurance_CBDC_JFMI_Figure_Table_Legends_v8.docx"
COVER_LETTER = OUT / "Composable_Assurance_CBDC_JFMI_Cover_Letter_TEMPLATE_v8.docx"
FIGURE_OUT = OUT / "editable_figures"

JFMI_TITLE = "Composable assurance for sovereign digital currency (CBDC): An evidence-gated qualification framework"
RUNNING_TITLE = "Composable CBDC assurance"

CITATION_LABELS = {
    1: "Illes, Kosse, and Wierts 2025",
    2: "BIS Innovation Hub 2023b",
    3: "BIS Innovation Hub 2024",
    4: "BIS Innovation Hub and Institute of International Finance 2026",
    5: "BIS Innovation Hub 2026",
    6: "Gray and Lamport 2006",
    7: "Castro and Liskov 1999",
    8: "Yin et al. 2019",
    9: "NIST 2025",
    10: "Financial Action Task Force 2026",
    11: "U.S. Treasury, Office of Foreign Assets Control 2026b",
    12: "U.S. Treasury, Office of Foreign Assets Control 2026c",
    13: "U.S. Treasury, Office of Foreign Assets Control 2026a",
    14: "Bidder, Jackson, and Rottner 2025",
    15: "European Central Bank 2025b",
    16: "European Central Bank 2025a",
    17: "European Central Bank 2024",
    18: "Banco Central do Brasil 2025",
    19: "BIS Innovation Hub 2023a",
    20: "BIS Innovation Hub 2025",
    21: "BIS Innovation Hub 2023c",
    22: "BIS Innovation Hub 2023d",
    23: "W3C 2026",
    24: "NIST 2024a",
    25: "NIST 2024b",
    26: "Financial Action Task Force 2020",
    27: "Financial Action Task Force 2023",
    28: "Cross-border Payments Interoperability and Extension Task Force 2026",
    29: "ISO n.d.",
    30: "CPMI-IOSCO 2012",
    31: "Basel Committee on Banking Supervision 2021",
    32: "European Union 2016",
    33: "Murphy et al. 2024",
    34: "Kao et al. 2025",
    35: "Lee et al. 2021",
    36: "Homoliak et al. 2023",
    37: "Reslow, Soderberg, and Tsuda 2024",
    38: "BIS Consultative Group on Innovation and the Digital Economy 2023",
    39: "Pocher and Veneris 2022",
    40: "Bharathan and Pillai 2022",
    41: "Michalopoulos et al. 2025",
    42: "Mullins et al. 2025",
    43: "Bernardo et al. 2025",
    44: "Jin and Xia 2022",
}

APA_REFERENCES = [
    "Banco Central do Brasil. (2025). Instant Payments System (SPI) annual report 2024.",
    "Basel Committee on Banking Supervision. (2021, March 31). Principles for operational resilience. https://www.bis.org/bcbs/publ/d516.htm",
    "Bidder, R., Jackson, T., & Rottner, M. (2025). CBDC and banks: Disintermediating fast and slow (BIS Working Papers No. 1280). Bank for International Settlements.",
    "BIS Consultative Group on Innovation and the Digital Economy. (2023, December 12). High-level technical requirements for a functional central bank digital currency architecture. https://www.bis.org/publ/othp82.htm",
    "BIS Innovation Hub. (2023a). Project Aurora: The power of data, technology and collaboration to combat money laundering.",
    "BIS Innovation Hub. (2023b). Project Icebreaker: Breaking new paths in cross-border retail CBDC payments.",
    "BIS Innovation Hub. (2023c). Project Polaris: A security and resilience framework for CBDC systems.",
    "BIS Innovation Hub. (2023d). Project Tourbillon: Exploring privacy, scalability and quantum-safe cryptography for CBDC.",
    "BIS Innovation Hub. (2024). Project Mandala: Shaping the future of cross-border payments compliance.",
    "BIS Innovation Hub. (2025). Project Hertha: Identifying financial crime patterns in payment systems.",
    "BIS Innovation Hub. (2026, January 29). Project FuSSE: Exploring flexible, scalable and secure settlement engines.",
    "BIS Innovation Hub, & Institute of International Finance. (2026, July 30). Project Agorá: Exploring tokenisation of wholesale cross-border payments. https://www.bis.org/about/bisih/topics/fmis/agora.htm",
    "Castro, M., & Liskov, B. (1999). Practical Byzantine fault tolerance. In Proceedings of OSDI (pp. 173-186).",
    "CPMI-IOSCO. (2012). Principles for financial market infrastructures.",
    "Cross-border Payments Interoperability and Extension Task Force. (2026). Enhancing cross-border payments: Fast payment system interlinking. Bank for International Settlements. https://www.bis.org/cpmi/pietf/fps_feb_2026.pdf",
    "European Central Bank. (2024, February 14). TIPS capacity elements (Requirement TIPS.UR.10.120). https://www.ecb.europa.eu/paym/target/target-professional-use-documents-links/tips/shared/pdf/tipsmeetdoc/ecb.tipsmeetdoc240214_TIPS-CG_TIPS_Capacity.en.pdf",
    "European Central Bank. (2025a, October 30). Preparation phase of a digital euro: Closing report. https://www.ecb.europa.eu/euro/digital_euro/progress/shared/pdf/ecb.deprp202510.en.pdf",
    "European Central Bank. (2025b, October 22). Technical data on the financial stability impact of the digital euro. https://www.ecb.europa.eu/euro/digital_euro/timeline/profuse/html/index.en.html",
    "European Union. (2016). General Data Protection Regulation, Regulation (EU) 2016/679.",
    "Financial Action Task Force. (2020, March 5). Guidance on digital identity. https://www.fatf-gafi.org/content/dam/fatf-gafi/guidance/Guidance-on-Digital-Identity.pdf.coredownload.pdf",
    "Financial Action Task Force. (2023, March 10). Guidance on beneficial ownership of legal persons. https://www.fatf-gafi.org/en/publications/Fatfrecommendations/Guidance-Beneficial-Ownership-Legal-Persons.html",
    "Financial Action Task Force. (2026). Methodology for assessing technical compliance with the FATF Recommendations and the effectiveness of AML/CFT/CPF systems. https://www.fatf-gafi.org/en/publications/Mutualevaluations/Fatf-methodology.html",
    "Gray, J., & Lamport, L. (2006). Consensus on transaction commit. ACM Transactions on Database Systems, 31(1), 133-160. https://doi.org/10.1145/1132863.1132867",
    "Homoliak, I., Perešíni, M., Holop, P., Handzuš, J., & Casino, F. (2023). CBDC-AquaSphere: Interoperable central bank digital currency built on trusted computing and blockchain. arXiv. https://doi.org/10.48550/arXiv.2305.16893",
    "Illes, A., Kosse, A., & Wierts, P. (2025). Advancing in tandem: Results of the 2024 BIS survey on central bank digital currencies and crypto (BIS Papers No. 159). Bank for International Settlements.",
    "ISO. (n.d.). ISO 20022: Universal financial industry message scheme. https://www.iso20022.org/",
    "Kao, K., Chen, K., Aldersey, B., Forte Walker, S., & Soana, G. (2025). Financial integrity implications of retail central bank digital currencies (IMF Fintech Note 2025/010). International Monetary Fund. https://doi.org/10.5089/9798229029308.063",
    "Lee, Y., Son, B., Jang, H., Byun, J., Yoon, T., & Lee, J. (2021). Atomic cross-chain settlement model for central banks digital currency. Information Sciences, 580, 838-856. https://doi.org/10.1016/j.ins.2021.09.040",
    "Murphy, K., et al. (2024). Central bank digital currency data use and privacy protection (IMF Fintech Note 2024/004). International Monetary Fund. https://doi.org/10.5089/9798400286971.063",
    "National Institute of Standards and Technology. (2024a). Module-lattice-based key-encapsulation mechanism standard (FIPS 203). https://doi.org/10.6028/NIST.FIPS.203",
    "National Institute of Standards and Technology. (2024b). Module-lattice-based digital signature standard (FIPS 204). https://doi.org/10.6028/NIST.FIPS.204",
    "National Institute of Standards and Technology. (2025). Guidelines for evaluating differential privacy guarantees (Special Publication 800-226). https://doi.org/10.6028/NIST.SP.800-226",
    "Pocher, N., & Veneris, A. (2022). Privacy and transparency in CBDCs: A regulation-by-design AML/CFT scheme. IEEE Transactions on Network and Service Management, 19(2), 1776-1788. https://doi.org/10.1109/TNSM.2021.3136984",
    "Reslow, A., Soderberg, G., & Tsuda, N. (2024). Cross-border payments with retail central bank digital currencies: Design and policy considerations (IMF Fintech Note 2024/002). International Monetary Fund. https://doi.org/10.5089/9798400272035.063",
    "U.S. Treasury, Office of Foreign Assets Control. (2026a). Entities owned by blocked persons: 50 percent rule (FAQs 398-402).",
    "U.S. Treasury, Office of Foreign Assets Control. (2026b). Frequently asked questions on Advanced Sanctions List Standard.",
    "U.S. Treasury, Office of Foreign Assets Control. (2026c). How Sanctions List Search works (FAQs 246-250).",
    "W3C. (2026, April 7). Data Integrity BBS Cryptosuites v1.0 (Candidate Recommendation Draft). https://www.w3.org/TR/vc-di-bbs/",
    "Yin, M., Malkhi, D., Reiter, M. K., Gueta, G. G., & Abraham, I. (2019). HotStuff: BFT consensus with linearity and responsiveness. In Proceedings of the ACM Symposium on Principles of Distributed Computing (pp. 347-356). https://doi.org/10.1145/3293611.3331591",
    "Bharathan, V., & Pillai, M. (2022). Central bank digital currency towards a composable standards-based implementation. SSRN. https://doi.org/10.2139/ssrn.4251613",
    "Bernardo, M., Calandra, F., Esposito, A., & Fabris, F. (2025). On the operational resilience of CBDC: Threats and prospects of formal validation for offline payments. arXiv. https://doi.org/10.48550/arXiv.2508.08064",
    "Michalopoulos, P., Olowookere, O., Pocher, N., Sedlmeir, J., Veneris, A., & Puri, P. (2025). Privacy and compliance design options in offline central bank digital currencies. IEEE Transactions on Network and Service Management, 22(5), 3748-3763. https://doi.org/10.1109/TNSM.2025.3575367",
    "Mullins, I., Brataniec, P., Jamroz, K., & Ade, G. (2025). Systems and methods for providing interoperability in a CBDC network (WIPO Patent No. WO2025085074A1). https://patents.google.com/patent/WO2025085074A1/en",
    "Jin, S. Y., & Xia, Y. (2022). CEV framework: A central bank digital currency evaluation and verification framework with a focus on consensus algorithms and operating architectures. IEEE Access, 10, 63698-63714. https://doi.org/10.1109/ACCESS.2022.3183092",
]

FIGURE_SOURCES = [
    "v7_assurance_stack.svg",
    "v7_architecture.svg",
    "v7_cross_border_sequence.svg",
    "v7_privacy_dp_pipeline.svg",
    "privacy_v7_learned_attack.svg",
    "dp_v7_utility.svg",
    "v7_aml_evaluation.svg",
    "aml_v7_families.svg",
    "v7_sanctions_workflow.svg",
    "sanctions_v7_policy_frontier.svg",
    "sanctions_v7_workload.svg",
    "v7_economic_decision.svg",
    "economic_v7_identified_set.svg",
    "v7_national_qualification.svg",
    "performance_v7_scenarios.svg",
    "performance_v7_capacity_bounds.svg",
]


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def convert_citations(text: str) -> str:
    def expand_range(match: re.Match[str]) -> str:
        start, end = int(match.group(1)), int(match.group(2))
        return "(" + "; ".join(CITATION_LABELS[number] for number in range(start, end + 1)) + ")"

    text = re.sub(r"\[(\d+)\]\s*[-–]\s*\[(\d+)\]", expand_range, text)
    text = re.sub(r"\[(\d+)\]", lambda match: f"({CITATION_LABELS[int(match.group(1))]})", text)
    # Join adjacent parenthetical citations produced from comma-separated numeric citations.
    while re.search(r"\)\s*,\s*\(", text):
        text = re.sub(r"\)\s*,\s*\(", "; ", text)
    return text


def set_page_footer(document: Document, text: str) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(text + " | ")
        run.font.size = Pt(7.1)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)


def remove_hard_page_breaks(document: Document) -> None:
    for element in list(document.element.body.iter(qn("w:br"))):
        if element.get(qn("w:type")) == "page":
            element.getparent().remove(element)


def insert_before(anchor, paragraph) -> None:
    anchor._p.addprevious(paragraph._p)


def count_words(document: Document) -> int:
    words: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "References":
            break
        words.extend(re.findall(r"\b[\w-]+\b", paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                words.extend(re.findall(r"\b[\w-]+\b", cell.text))
    return len(words)


def numbered_caption_count(document: Document, prefix: str) -> int:
    pattern = re.compile(rf"^{prefix}\s+(\d+)\.")
    found = set()
    for paragraph in document.paragraphs:
        match = pattern.match(paragraph.text.strip())
        if match:
            found.add(int(match.group(1)))
    return len(found)


def style_document(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def build_main() -> tuple[Document, int, int, int]:
    document = Document(SOURCE)
    remove_hard_page_breaks(document)
    document.paragraphs[0].text = JFMI_TITLE
    document.paragraphs[0].style = "Title"
    for paragraph in list(document.paragraphs[1:4]):
        remove_paragraph(paragraph)

    for paragraph in document.paragraphs:
        converted = convert_citations(paragraph.text)
        if converted != paragraph.text:
            replace_paragraph_text(paragraph, converted)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    converted = convert_citations(paragraph.text)
                    if converted != paragraph.text:
                        replace_paragraph_text(paragraph, converted)

    references_index = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "References")
    references = document.paragraphs[references_index]
    for paragraph in list(document.paragraphs[references_index + 1 :]):
        remove_paragraph(paragraph)

    for heading_text, body_text in [
        (
            "Acknowledgements",
            "[AI-USE ACKNOWLEDGEMENT NOT APPROVED BY THE ACCOUNTABLE AUTHOR.] OpenAI Codex was used to organize background sources and reproducibility artifacts, assist with code and document editing, and improve language. The accountable author must verify the final manuscript, analyses, citations, code, and disclosure before submission.",
        ),
        (
            "Declarations of Interest",
            "Funding: Self-funded. Conflicts of interest: [ACCOUNTABLE AUTHOR CONFIRMATION REQUIRED BEFORE SUBMISSION. No conflict-of-interest declaration is asserted by this provisional build.]",
        ),
    ]:
        heading = document.add_paragraph(heading_text, style="Heading 1")
        insert_before(references, heading)
        body = document.add_paragraph(body_text, style="Body Text")
        insert_before(references, body)

    for reference in sorted(APA_REFERENCES, key=str.casefold):
        paragraph = document.add_paragraph(reference)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.94
        for run in paragraph.runs:
            run.font.size = Pt(7.8)

    document.core_properties.title = JFMI_TITLE
    document.core_properties.subject = "Provisional blinded JFMI manuscript"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = (
        "Blinded provisional adaptation; accountable author confirmation remains required for affiliation, "
        "postal address, corresponding-author designation, conflicts, AI-use disclosure, repository deposit and submission."
    )
    set_page_footer(document, "Blinded manuscript")
    style_document(document)

    word_count = count_words(document)
    figure_count = numbered_caption_count(document, "Figure")
    table_count = numbered_caption_count(document, "Table")
    document.save(MAIN)
    return document, word_count, figure_count, table_count


def build_title_page(word_count: int, figure_count: int, table_count: int) -> None:
    document = Document()
    style_document(document)
    title = document.add_paragraph(JFMI_TITLE)
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice = document.add_paragraph("PROVISIONAL TITLE-PAGE TEMPLATE — ACCOUNTABLE AUTHOR CONFIRMATION REQUIRED")
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.runs[0].bold = True
    notice.runs[0].font.color.rgb = RGBColor(156, 0, 6)

    rows = [
        ("Journal", "Journal of Financial Market Infrastructures (provisional target)"),
        ("Article type", "Research paper"),
        ("Author 1", "Zubaer Mahmood Zubraj"),
        ("Affiliation", "[ACCOUNTABLE AUTHOR TO CONFIRM]"),
        ("Postal address", "[ACCOUNTABLE AUTHOR TO CONFIRM]"),
        ("Email", "zmzubraj@gmail.com"),
        ("ORCID", "https://orcid.org/0009-0004-3251-3385"),
        ("Corresponding author", "[ACCOUNTABLE AUTHOR TO CONFIRM]"),
        ("Running title", RUNNING_TITLE),
        ("Word count", f"{word_count:,} words, including abstract/captions/tables and excluding references"),
        ("Figures", str(figure_count)),
        ("Tables", str(table_count)),
        ("Funding", "Self-funded"),
        ("Conflicts of interest", "[ACCOUNTABLE AUTHOR TO CONFIRM]"),
        ("Repository/DOI/license", "MIT code licence confirmed; [PUBLIC REPOSITORY URL AND DOI NOT YET ASSIGNED — AUTHOR RELEASE ACTION REQUIRED]"),
        ("Submission exclusivity", "Confirmed by the author on 29 August 2026: not published or under consideration elsewhere"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True

    document.add_heading("AI-use acknowledgement — not approved", level=1)
    document.add_paragraph(
        "OpenAI Codex was used to organize background sources and reproducibility artifacts, "
        "assist with code and document editing, and improve language. The accountable author "
        "must verify the final manuscript, analyses, citations, code, and disclosure. "
        "[THE AUTHOR ANSWERED NO TO APPROVAL ON 29 AUGUST 2026; SUBMISSION REMAINS BLOCKED UNTIL A TRUTHFUL DISCLOSURE IS APPROVED.]"
    )
    document.add_heading("Required author action", level=1)
    document.add_paragraph(
        "Replace every remaining bracketed field, confirm the conflict declaration and a truthful AI-use statement, and verify "
        "that the title-page identity exactly matches the submission-portal account."
    )
    document.core_properties.title = "JFMI title-page template"
    document.core_properties.author = "Zubaer Mahmood Zubraj"
    document.save(TITLE_PAGE)


def build_legends(document: Document) -> None:
    legends = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^(Figure|Table)\s+\d+\.", text):
            legends.append(text)
    legends.sort(key=lambda value: (0 if value.startswith("Figure") else 1, int(re.search(r"\d+", value).group())))

    output = Document()
    style_document(output)
    output.add_heading("Figure and table legends", level=1)
    output.add_paragraph(
        "Provisional companion file. The same numbered figures, tables, and captions remain embedded in the blinded manuscript."
    )
    for legend in legends:
        paragraph = output.add_paragraph(legend)
        paragraph.paragraph_format.space_after = Pt(5)
    output.core_properties.author = ""
    output.core_properties.last_modified_by = ""
    output.save(LEGENDS)


def build_cover_letter() -> None:
    document = Document()
    style_document(document)
    notice = document.add_paragraph("PROVISIONAL OPTIONAL COVER LETTER — DO NOT SUBMIT WITH OPEN GATES")
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.runs[0].bold = True
    notice.runs[0].font.color.rgb = RGBColor(156, 0, 6)

    document.add_paragraph("29 August 2026")
    document.add_paragraph("Editor-in-Chief\nThe Journal of Financial Market Infrastructures")
    document.add_paragraph("Dear Editor-in-Chief,")
    document.add_paragraph(
        f'I submit the manuscript “{JFMI_TITLE}” for provisional consideration as a research paper in '
        "The Journal of Financial Market Infrastructures."
    )
    document.add_paragraph(
        "The paper develops an evidence-gated qualification framework for sovereign CBDC payment infrastructure. "
        "It binds monetary invariants, cross-border settlement evidence, privacy and financial-integrity evaluation, "
        "policy limits, and operational-scale claims to explicit maturity ceilings. Its intended contribution is the "
        "composable qualification method rather than priority for each underlying component."
    )
    document.add_paragraph(
        "The manuscript is relevant to the journal's coverage of payment and settlement systems, settlement risk and "
        "interdependencies, non-bank payment service providers, retail payment infrastructure, oversight, and "
        "infrastructure-related standardization. The evidence comprises formal modelling, reproducible synthetic and "
        "public-list benchmarks, an internal local prototype, and a queueing digital twin. It does not claim certified-HSM, "
        "physical multi-region, institutional-field, national-deployment, or independent-replication evidence."
    )
    document.add_paragraph(
        "The author confirms that the manuscript has not been published and is not under consideration or in press "
        "elsewhere. The study is a non-human computational systems study using public and synthetic materials; no human "
        "participant, clinical intervention, or identifiable personal-data result is claimed. Funding is self-funded."
    )
    gate = document.add_paragraph(
        "[ACCOUNTABLE AUTHOR MUST INSERT AND APPROVE THE CONFLICT-OF-INTEREST STATEMENT AND A JOURNAL-COMPLIANT, "
        "TRUTHFUL AI-USE ACKNOWLEDGEMENT BEFORE THIS LETTER MAY BE SUBMITTED.]"
    )
    gate.runs[0].bold = True
    gate.runs[0].font.color.rgb = RGBColor(156, 0, 6)
    document.add_paragraph(
        "The blinded manuscript excludes author identity. A separate title page provides the confirmed author name, "
        "email, ORCID, funding statement, and exclusivity confirmation; affiliation, postal address, and the "
        "corresponding-author designation remain open author fields."
    )
    document.add_paragraph("Thank you for considering the manuscript.")
    document.add_paragraph(
        "Sincerely,\nZubaer Mahmood Zubraj\nEmail: zmzubraj@gmail.com\n"
        "ORCID: https://orcid.org/0009-0004-3251-3385\n"
        "[AFFILIATION, POSTAL ADDRESS, AND CORRESPONDING-AUTHOR DESIGNATION TO BE CONFIRMED]"
    )
    document.core_properties.title = "Provisional optional JFMI cover-letter template"
    document.core_properties.subject = "Not submission-ready while bracketed author gates remain"
    document.core_properties.author = "Zubaer Mahmood Zubraj"
    document.core_properties.comments = "Optional local template; the official JFMI guidelines checked on 29 August 2026 do not list a cover letter as mandatory."
    document.save(COVER_LETTER)


def copy_vector_figures() -> None:
    FIGURE_OUT.mkdir(parents=True, exist_ok=True)
    for number, name in enumerate(FIGURE_SOURCES, start=1):
        source = ROOT / "figures" / name
        if not source.exists():
            raise FileNotFoundError(f"missing vector figure source: {source}")
        shutil.copy2(source, FIGURE_OUT / f"Figure_{number:02d}.svg")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    OUT.mkdir(parents=True, exist_ok=True)
    document, word_count, figure_count, table_count = build_main()
    if figure_count != 16 or table_count != 22:
        raise RuntimeError(f"unexpected caption counts: figures={figure_count}, tables={table_count}")
    build_title_page(word_count, figure_count, table_count)
    build_legends(document)
    build_cover_letter()
    copy_vector_figures()
    print(MAIN)
    print(TITLE_PAGE)
    print(LEGENDS)
    print(COVER_LETTER)
    print(f"word_count={word_count} figures={figure_count} tables={table_count}")


if __name__ == "__main__":
    main()
