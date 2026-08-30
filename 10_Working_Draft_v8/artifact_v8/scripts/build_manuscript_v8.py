from __future__ import annotations

import runpy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
SOURCE = OUTPUT / "Composable_Assurance_CBDC_National_Qualification_Manuscript_v7.docx"
TARGET = OUTPUT / "Composable_Assurance_CBDC_Author_Draft_v8.docx"
PUBLISHING_TITLE = "Composable assurance for sovereign digital currency (CBDC): An evidence-gated qualification framework"

ABSTRACT = (
    "A central bank digital currency (CBDC) architecture must jointly preserve monetary "
    "singleness, privacy, financial integrity, legal finality, and operational resilience. "
    "This study presents an evidence-constrained composition framework for reserve-settled "
    "domestic conversion and payment-versus-payment across independently governed sovereign "
    "ledgers. A formal state model explores 2,013 states and 19,480 transitions under a declared "
    "two-key equivocation bound. A local eleven-service prototype binds bilateral compliance "
    "authorizations, unequal-currency quotes, and final prepare receipts to 5-of-7 quorum "
    "certificates, while testing restart, duplicate delivery, stale evidence, and conflicting "
    "votes. Synthetic privacy experiments show that shielding and batching reduce network-observer "
    "linkability under generator shift, although residual leakage remains. A registered "
    "person-level differential-privacy service enforces a twelve-release epsilon budget. Across "
    "sixteen synthetic AML graphs, cross-institution features improve mean PR-AUC by 0.058, with "
    "heterogeneity across graph families. Calibrated sanctions screening exposes the operational "
    "trade-off between recall and false-alert workload. Economic analysis replaces a universal "
    "holding limit with a partially identified adaptive rule. A queueing digital twin passes "
    "declared 2,000-20,000 TPS scenarios but fails an overload control. The contribution is a "
    "reproducible qualification pathway, not evidence of production or national deployment."
)

KEYWORDS = (
    "Keywords: central bank digital currency; cross-border settlement; differential privacy; "
    "AML/CFT; operational resilience; reproducibility"
)

KEY_MESSAGES = [
    "CBDC assurance must be evaluated as a composed system, not isolated modules.",
    "Cross-border decisions bind monetary, FX, compliance, and prepare evidence.",
    "Privacy, AML, sanctions, and policy claims remain bounded by their evidence.",
    "Scale qualification must progress from models to physical and field evidence.",
]

RELATED_WORK_ORIGINAL = (
    "Project Icebreaker demonstrated cross-border retail settlement through two domestic CBDC "
    "legs funded by FX providers [2]. Project Mandala developed rules and proof engines for "
    "compliance-by-design [3]. Project Agorá demonstrated atomic multi-currency wholesale "
    "settlement on a shared programmable platform [4]. Project FuSSE reported 10,000 "
    "transactions per second in a controlled modular proof of concept while explicitly identifying "
    "scaling and operational challenges [5]. These projects provide essential components, but they "
    "use different market structures and evidence scopes."
)

RELATED_WORK_UPDATED = (
    "Project Icebreaker demonstrated cross-border retail settlement through two domestic CBDC "
    "legs funded by FX providers [2]. Project Mandala developed rules and proof engines for "
    "compliance-by-design [3]. Project Agorá demonstrated atomic multi-currency wholesale "
    "settlement on a shared programmable platform and, in July 2026, completed 17 controlled "
    "real-value scenarios involving 28 financial institutions and central banks and approximately "
    "CHF 800,000; the BIS describes this as prototype testing rather than a finished product [4]. "
    "Project FuSSE reported 10,000 transactions per second in a controlled modular proof of "
    "concept while explicitly identifying scaling and operational challenges [5]. These projects "
    "provide essential components, but they use different market structures and evidence scopes. "
    "Academic and policy work separately addresses atomic cross-chain CBDC settlement, "
    "TEE-based interoperability across independent semi-centralized ledgers, five-element "
    "cross-border design choices, functional CBDC architecture requirements, and "
    "privacy/AML regulation-by-design [35]-[39]. A standards-based composable CBDC implementation "
    "already links token and contract networks [40]; later work extends offline privacy/compliance "
    "classification [41], architecture-agnostic cross-system checkpoints [42], and formal-method "
    "arguments for CBDC operational resilience [43]. Jin and Xia's CEV framework already "
    "recommends CBDC consensus and operating architectures and validates proposals through "
    "experiments and formal analysis [44]. The claimed difference here is therefore not priority "
    "for any isolated component or for CBDC evaluation and verification as a broad category. It is "
    "the narrower joint claim-to-evidence qualification method "
    "that connects settlement safety, privacy leakage and differential privacy, AML and sanctions "
    "evaluation, adaptive holding limits, and falsifiable scale gates under one explicit maturity boundary."
)

REFERENCE_REPLACEMENTS = {
    '[4] BIS Innovation Hub and Institute of International Finance, “Project Agorá: a shared programmable platform for tokenised wholesale cross-border payments,” 27 May 2026.':
        '[4] BIS Innovation Hub and Institute of International Finance, “Project Agorá: exploring tokenisation of wholesale cross-border payments,” updated 30 Jul. 2026, https://www.bis.org/about/bisih/topics/fmis/agora.htm (accessed 29 Aug. 2026).',
    '[9] NIST, “Guidelines for Evaluating Differential Privacy Guarantees,” Special Publication 800-226, Mar. 2025.':
        '[9] NIST, “Guidelines for Evaluating Differential Privacy Guarantees,” Special Publication 800-226, Mar. 2025, doi:10.6028/NIST.SP.800-226.',
    '[10] Financial Action Task Force, “Methodology for Assessing Technical Compliance with the FATF Recommendations and the Effectiveness of AML/CFT/CPF Systems,” updated June 2026.':
        '[10] Financial Action Task Force, “Methodology for Assessing Technical Compliance with the FATF Recommendations and the Effectiveness of AML/CFT/CPF Systems,” updated Jun. 2026, https://www.fatf-gafi.org/en/publications/Mutualevaluations/Fatf-methodology.html (accessed 29 Aug. 2026).',
    '[15] European Central Bank, “Technical data on the financial stability impact of the digital euro,” seminar technical data, 22 Oct. 2025.':
        '[15] European Central Bank, “Technical data on the financial stability impact of the digital euro,” technical annex, 22 Oct. 2025, https://www.ecb.europa.eu/euro/digital_euro/timeline/profuse/html/index.en.html (accessed 29 Aug. 2026).',
    '[16] European Central Bank, “Preparation phase of a digital euro: closing report,” 30 Oct. 2025.':
        '[16] European Central Bank, “Preparation phase of a digital euro: closing report,” 30 Oct. 2025, https://www.ecb.europa.eu/euro/digital_euro/progress/shared/pdf/ecb.deprp202510.en.pdf (accessed 29 Aug. 2026).',
    '[17] European Central Bank, “TIPS Capacity Elements,” requirement TIPS.UR.10.120, Feb. 2024; capacity reconfirmed Feb. 2026.':
        '[17] European Central Bank, “TIPS Capacity Elements,” requirement TIPS.UR.10.120, 14 Feb. 2024, https://www.ecb.europa.eu/paym/target/target-professional-use-documents-links/tips/shared/pdf/tipsmeetdoc/ecb.tipsmeetdoc240214_TIPS-CG_TIPS_Capacity.en.pdf (accessed 29 Aug. 2026).',
    '[23] W3C, “Data Integrity BBS Cryptosuites v1.0,” Candidate Recommendation Draft, accessed 6 Aug. 2026.':
        '[23] W3C, “Data Integrity BBS Cryptosuites v1.0,” Candidate Recommendation Draft, 7 Apr. 2026, work in progress, https://www.w3.org/TR/vc-di-bbs/ (accessed 29 Aug. 2026).',
    '[24] NIST, “Module-Lattice-Based Key-Encapsulation Mechanism Standard,” FIPS 203, 2024.':
        '[24] NIST, “Module-Lattice-Based Key-Encapsulation Mechanism Standard,” FIPS 203, 13 Aug. 2024, doi:10.6028/NIST.FIPS.203.',
    '[25] NIST, “Module-Lattice-Based Digital Signature Standard,” FIPS 204, 2024.':
        '[25] NIST, “Module-Lattice-Based Digital Signature Standard,” FIPS 204, 13 Aug. 2024, doi:10.6028/NIST.FIPS.204.',
    '[26] Financial Action Task Force, “Guidance on Digital Identity,” 6 Mar. 2020.':
        '[26] Financial Action Task Force, “Guidance on Digital Identity,” 5 Mar. 2020, https://www.fatf-gafi.org/content/dam/fatf-gafi/guidance/Guidance-on-Digital-Identity.pdf.coredownload.pdf (accessed 29 Aug. 2026).',
    '[27] Financial Action Task Force, “Guidance on Beneficial Ownership of Legal Persons,” Mar. 2023.':
        '[27] Financial Action Task Force, “Guidance on Beneficial Ownership of Legal Persons,” 10 Mar. 2023, https://www.fatf-gafi.org/en/publications/Fatfrecommendations/Guidance-Beneficial-Ownership-Legal-Persons.html (accessed 29 Aug. 2026).',
    '[28] CPMI, “Fast payment system interlinking,” updated report, Feb. 2026.':
        '[28] Cross-border Payments Interoperability and Extension Task Force, “Enhancing cross-border payments: fast payment system interlinking,” Feb. 2026, https://www.bis.org/cpmi/pietf/fps_feb_2026.pdf (accessed 29 Aug. 2026).',
    '[31] Basel Committee on Banking Supervision, “Principles for Operational Resilience,” Mar. 2021.':
        '[31] Basel Committee on Banking Supervision, “Principles for operational resilience,” 31 Mar. 2021, https://www.bis.org/bcbs/publ/d516.htm (accessed 29 Aug. 2026).',
    '[33] International Monetary Fund, “Central Bank Digital Currency Data Use and Privacy Protection,” Fintech Note, 2024.':
        '[33] K. Murphy et al., “Central Bank Digital Currency Data Use and Privacy Protection,” IMF Fintech Note 2024/004, 30 Aug. 2024, doi:10.5089/9798400286971.063.',
    '[34] International Monetary Fund, “Central Bank Digital Currencies and Financial Integrity,” Fintech Note, 2025.':
        '[34] K. Kao, K. Chen, B. Aldersey, S. Forte Walker and G. Soana, “Financial Integrity Implications of Retail Central Bank Digital Currencies,” IMF Fintech Note 2025/010, 17 Nov. 2025, doi:10.5089/9798229029308.063.',
}

ADDITIONAL_REFERENCES = [
    '[35] Y. Lee, B. Son, H. Jang, J. Byun, T. Yoon and J. Lee, “Atomic cross-chain settlement model for central banks digital currency,” Information Sciences, vol. 580, pp. 838-856, 2021, doi:10.1016/j.ins.2021.09.040.',
    '[36] I. Homoliak, M. Perešíni, P. Holop, J. Handzuš and F. Casino, “CBDC-AquaSphere: Interoperable Central Bank Digital Currency Built on Trusted Computing and Blockchain,” arXiv:2305.16893, 2023, doi:10.48550/arXiv.2305.16893.',
    '[37] A. Reslow, G. Soderberg and N. Tsuda, “Cross-Border Payments with Retail Central Bank Digital Currencies: Design and Policy Considerations,” IMF Fintech Note 2024/002, 15 May 2024, doi:10.5089/9798400272035.063.',
    '[38] BIS Consultative Group on Innovation and the Digital Economy, “High-level technical requirements for a functional central bank digital currency architecture,” 12 Dec. 2023, https://www.bis.org/publ/othp82.htm (accessed 29 Aug. 2026).',
    '[39] N. Pocher and A. Veneris, “Privacy and Transparency in CBDCs: A Regulation-by-Design AML/CFT Scheme,” IEEE Transactions on Network and Service Management, vol. 19, no. 2, pp. 1776-1788, 2022, doi:10.1109/TNSM.2021.3136984.',
    '[40] V. Bharathan and M. Pillai, “Central Bank Digital Currency Towards A Composable Standards-Based Implementation,” SSRN, 18 Oct. 2022, doi:10.2139/ssrn.4251613.',
    '[41] P. Michalopoulos, O. Olowookere, N. Pocher, J. Sedlmeir, A. Veneris and P. Puri, “Privacy and Compliance Design Options in Offline Central Bank Digital Currencies,” IEEE Transactions on Network and Service Management, vol. 22, no. 5, pp. 3748-3763, 2025, doi:10.1109/TNSM.2025.3575367.',
    '[42] I. Mullins, P. Brataniec, K. Jamroz and G. Ade, “Systems and methods for providing interoperability in a CBDC network,” WIPO Patent WO2025085074A1, published 24 Apr. 2025, priority 17 Oct. 2023, https://patents.google.com/patent/WO2025085074A1/en.',
    '[43] M. Bernardo, F. Calandra, A. Esposito and F. Fabris, “On the Operational Resilience of CBDC: Threats and Prospects of Formal Validation for Offline Payments,” arXiv:2508.08064, 2025, doi:10.48550/arXiv.2508.08064.',
    '[44] S. Y. Jin and Y. Xia, “CEV Framework: A Central Bank Digital Currency Evaluation and Verification Framework With a Focus on Consensus Algorithms and Operating Architectures,” IEEE Access, vol. 10, pp. 63698-63714, 2022, doi:10.1109/ACCESS.2022.3183092.',
]


def find_paragraph(document: Document, exact_text: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == exact_text:
            return index
    raise RuntimeError(f"paragraph not found: {exact_text}")


def format_title_page(document: Document) -> None:
    """Apply explicit title-page typography instead of inheriting the v7 run styles."""
    title = document.paragraphs[0]
    title.clear()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.line_spacing = 1.05
    title.paragraph_format.keep_with_next = True
    title_line_1 = title.add_run(
        "Composable assurance for sovereign digital currency"
    )
    title_line_1.add_break()
    title_line_2 = title.add_run("(CBDC): An evidence-gated qualification framework")
    for title_run in (title_line_1, title_line_2):
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(31, 78, 121)

    spacer = document.paragraphs[1]
    spacer.text = ""
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)

    author = document.paragraphs[2]
    author.text = "Zubaer Mahmood Zubraj"
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(0)
    author.paragraph_format.space_after = Pt(6)
    author.paragraph_format.keep_with_next = True
    author_run = author.runs[0]
    author_run.bold = True
    author_run.font.name = "Arial"
    author_run.font.size = Pt(12.5)

    metadata = document.paragraphs[3]
    metadata.clear()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(20)
    metadata.paragraph_format.line_spacing = 1.15
    metadata.paragraph_format.keep_with_next = True
    email_run = metadata.add_run("Email: zmzubraj@gmail.com")
    email_run.add_break()
    orcid_run = metadata.add_run("ORCID: https://orcid.org/0009-0004-3251-3385")
    for metadata_run in (email_run, orcid_run):
        metadata_run.font.name = "Arial"
        metadata_run.font.size = Pt(10)


def build() -> Path:
    runpy.run_path(str(ROOT / "scripts" / "build_manuscript_v7.py"), run_name="__main__")
    document = Document(SOURCE)

    format_title_page(document)

    abstract_index = find_paragraph(document, "Abstract")
    document.paragraphs[abstract_index + 1].text = ABSTRACT
    document.paragraphs[abstract_index + 2].text = KEYWORDS

    document.paragraphs[find_paragraph(document, RELATED_WORK_ORIGINAL)].text = RELATED_WORK_UPDATED
    for original, replacement in REFERENCE_REPLACEMENTS.items():
        document.paragraphs[find_paragraph(document, original)].text = replacement

    for reference in ADDITIONAL_REFERENCES:
        document.add_paragraph(reference)

    references_index = find_paragraph(document, "References")
    for paragraph in document.paragraphs[references_index + 1:]:
        for run in paragraph.runs:
            run.font.size = Pt(7.9)
        paragraph.paragraph_format.line_spacing = 0.96
        paragraph.paragraph_format.space_after = Pt(0.25)

    introduction_index = find_paragraph(document, "1. Introduction")
    anchor = document.paragraphs[introduction_index]._p
    heading = document.add_paragraph("Key messages", style="Heading 1")
    anchor.addprevious(heading._p)
    for message in KEY_MESSAGES:
        paragraph = document.add_paragraph(message, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        anchor.addprevious(paragraph._p)

    document.core_properties.title = PUBLISHING_TITLE
    document.core_properties.author = "Zubaer Mahmood Zubraj"
    document.core_properties.subject = "Author draft v8; evidence-constrained CBDC assurance"
    document.core_properties.keywords = (
        "central bank digital currency, cross-border settlement, differential privacy, "
        "AML/CFT, operational resilience, reproducibility"
    )
    document.core_properties.comments = (
        "Research evidence is formal, synthetic, laboratory-prototype, official-list snapshot, "
        "and queueing-simulation evidence; no production or national deployment claim."
    )
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = 1
        header_run = header.add_run("Composable CBDC assurance")
        header_run.font.name = "Arial"
        header_run.font.size = Pt(8)
        header_run.italic = True
        header_run.font.color.rgb = RGBColor(89, 89, 89)
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = 1
        run = footer.add_run("Zubaer Mahmood Zubraj | 29 August 2026 | ")
        run.font.size = Pt(7.1)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)
    document.save(TARGET)
    print(TARGET)
    return TARGET


if __name__ == "__main__":
    build()
