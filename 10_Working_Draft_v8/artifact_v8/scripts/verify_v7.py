from pathlib import Path
import json,re,zipfile,hashlib
import pandas as pd
from docx import Document
ROOT=Path(__file__).resolve().parents[1]; RES=ROOT/'results'; OUT=ROOT/'output'; DOC=OUT/'Composable_Assurance_CBDC_National_Qualification_Manuscript_v7.docx'; PDF=OUT/'Composable_Assurance_CBDC_National_Qualification_Manuscript_v7.pdf'
assert DOC.exists() and DOC.stat().st_size>100_000
assert PDF.exists() and PDF.stat().st_size>100_000
model=json.loads((RES/'cross_border_model_v5.json').read_text()); assert model['reachable_states']==2013 and model['transition_edges']==19480
p=json.loads((RES/'privacy_v7_experiment_summary.json').read_text()); assert p['candidate_sizes']==[100,1000,3000]
dp=json.loads((RES/'dp_v7_summary.json').read_text()); assert dp['l1_sensitivity']==3 and dp['observed_max_neighbor_l1_over_10000_tests']==3
aml=json.loads((RES/'aml_v7_summary.json').read_text()); assert aml['independent_graphs']==16 and aml['generator_families']==4 and aml['positive_effect_graphs']==12
san=json.loads((RES/'sanctions_v7_summary.json').read_text()); assert san['official_records_indexed']==19178 and san['official_aliases_indexed']==20010 and san['ownership_monotonicity_properties_passed']==500
econ=json.loads((RES/'economic_v7_summary.json').read_text()); assert econ['universal_optimum_identified'] is False
perf=json.loads((RES/'performance_v7_summary.json').read_text()); assert perf['national_scale_demonstrated'] is False and perf['scenarios']==7
scen=pd.read_csv(RES/'performance_v7_scenarios.csv'); assert (~scen.domestic_slo_pass).sum()==1
# manuscript scan
D=Document(DOC); text='\n'.join(p.text for p in D.paragraphs)
for t in D.tables:
    text+='\n'+'\n'.join(' | '.join(c.text for c in row.cells) for row in t.rows)
for forbidden in ['Smart Taka','ReasonedTransfer','ERC-20','ERC-621']:
    assert forbidden.lower() not in text.lower(), forbidden
body=text.split('References')[0]
# accept ranges as citations: build numbers from every [..] citation group
cited=set()
for a,b in re.findall(r'\[(\d+)\]\s*[-–]\s*\[(\d+)\]',body): cited.update(range(int(a),int(b)+1))
for grp in re.findall(r'\[([0-9,\-\s]+)\]',body):
    for part in grp.split(','):
        part=part.strip()
        if '-' in part:
            a,b=part.split('-',1); cited.update(range(int(a),int(b)+1))
        elif part.isdigit(): cited.add(int(part))
missing=[i for i in range(1,35) if i not in cited]
assert not missing, f'uncited references: {missing}'
# key text/result consistency
assert '8.43' in text and '3,571' in text and '1,643' in text
report={
 'status':'PASS','docx_bytes':DOC.stat().st_size,'pdf_bytes':PDF.stat().st_size,
 'references_cited':sorted(cited),'forbidden_terms_absent':True,
 'cross_border_states':model['reachable_states'],'aml_graphs':aml['independent_graphs'],
 'dp_sensitivity':dp['l1_sensitivity'],'sanctions_records':san['official_records_indexed'],
 'performance_overload_failures':int((~scen.domestic_slo_pass).sum())
}
(ROOT/'docs'/'VERIFICATION_V7.json').write_text(json.dumps(report,indent=2))
print(json.dumps(report,indent=2))
