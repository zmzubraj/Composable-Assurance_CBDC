from __future__ import annotations

import json
import re
import subprocess
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document

from build_jfmi_submission_v8 import APA_REFERENCES, CITATION_LABELS, JFMI_TITLE


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "jfmi_provisional"
DOCX = OUT / "Composable_Assurance_CBDC_JFMI_Blinded_Manuscript_v8.docx"
PDF = OUT / "Composable_Assurance_CBDC_JFMI_Blinded_Manuscript_v8.pdf"
TITLE_PAGE = OUT / "Composable_Assurance_CBDC_JFMI_Title_Page_TEMPLATE_v8.docx"
LEGENDS = OUT / "Composable_Assurance_CBDC_JFMI_Figure_Table_Legends_v8.docx"
COVER_LETTER = OUT / "Composable_Assurance_CBDC_JFMI_Cover_Letter_TEMPLATE_v8.docx"
COVER_LETTER_PDF = OUT / "Composable_Assurance_CBDC_JFMI_Cover_Letter_TEMPLATE_v8.pdf"
FIGURES = OUT / "editable_figures"
REPORT = ROOT / "docs" / "JFMI_VERIFICATION_V8.json"
INVENTORY = ROOT / "docs" / "JFMI_SUBMISSION_INVENTORY_V8.json"


def require(condition: bool, message: str) -> None:
    if not condition:
        raise RuntimeError(message)


def collect_text(document: Document) -> str:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    return text


def caption_numbers(document: Document, prefix: str) -> set[int]:
    pattern = re.compile(rf"^{prefix}\s+(\d+)\.")
    values = set()
    for paragraph in document.paragraphs:
        match = pattern.match(paragraph.text.strip())
        if match:
            values.add(int(match.group(1)))
    return values


def main() -> None:
    for path in (DOCX, PDF, TITLE_PAGE, LEGENDS, COVER_LETTER, COVER_LETTER_PDF, INVENTORY):
        require(path.is_file() and path.stat().st_size > 1_000, f"missing or unexpectedly small JFMI artifact: {path}")

    document = Document(DOCX)
    text = collect_text(document)
    lower = text.lower()
    require(document.core_properties.author in (None, ""), "blinded DOCX author metadata is not empty")
    require(document.core_properties.last_modified_by in (None, ""), "blinded DOCX last-modified-by is not empty")
    for forbidden in (
        "md zubaer mahmood zubraj",
        "zubraj14@gmail.com",
        "zubaer mahmood zubraj",
        "zmzubraj@gmail.com",
        "0009-0004-3251-3385",
        "independent researcher",
        "systems/cryptography reviewer",
        "privacy/aml-method reviewer",
    ):
        require(forbidden not in lower, f"forbidden identity/reviewer text in blinded manuscript: {forbidden}")
    require("@" not in text, "email-like identifier remains in blinded manuscript")

    title = document.paragraphs[0].text.strip()
    require(title == JFMI_TITLE, "JFMI publishing title changed")
    require(len(re.findall(r"\b[\w-]+\b", title)) <= 18, "JFMI title is not concise")
    abstract_index = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "Abstract")
    abstract = document.paragraphs[abstract_index + 1].text
    abstract_words = len(re.findall(r"\b[\w-]+\b", abstract))
    require(150 <= abstract_words <= 200, f"abstract outside JFMI 150-200 word range: {abstract_words}")
    require(not re.search(r"\([A-Za-z][^)]*\b(?:19|20)\d{2}[a-z]?\)", abstract), "abstract contains a citation")

    keyword_line = next(paragraph.text for paragraph in document.paragraphs if paragraph.text.startswith("Keywords:"))
    keyword_count = len(keyword_line.removeprefix("Keywords:").split(";"))
    require(4 <= keyword_count <= 6, f"keyword count outside JFMI 4-6 range: {keyword_count}")

    key_index = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "Key messages")
    key_messages = []
    for paragraph in document.paragraphs[key_index + 1 :]:
        if paragraph.style.name == "List Bullet":
            key_messages.append(paragraph.text.strip())
        elif key_messages:
            break
    require(3 <= len(key_messages) <= 4, f"JFMI requires 3-4 key messages, found {len(key_messages)}")
    require(all(len(message) <= 85 for message in key_messages), "a JFMI key message exceeds 85 characters")

    require("Acknowledgements" in text, "acknowledgements section missing")
    require("Declarations of Interest" in text, "declarations-of-interest section missing")
    require("OpenAI Codex" in text, "AI-use acknowledgement missing")
    require("ACCOUNTABLE AUTHOR CONFIRMATION REQUIRED" in text, "author declaration gate missing")

    references_heading_index = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "References")
    paragraph_body = "\n".join(paragraph.text for paragraph in document.paragraphs[:references_heading_index])
    table_body = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    body = paragraph_body + "\n" + table_body
    require(not re.search(r"\[\d+\](?:\s*[-–]\s*\[\d+\])?", body), "numeric citation remains in JFMI body")
    missing_labels = [number for number, label in CITATION_LABELS.items() if label not in body]
    require(not missing_labels, f"author-date citations missing from body: {missing_labels}")
    references = [paragraph.text.strip() for paragraph in document.paragraphs[references_heading_index + 1 :] if paragraph.text.strip()]
    require(len(references) == 44, f"expected 44 APA references, found {len(references)}")
    require(references == sorted(APA_REFERENCES, key=str.casefold), "reference list differs from the curated APA-order manifest")
    require(all(not reference.startswith("[") for reference in references), "numbered reference remains")

    figures = caption_numbers(document, "Figure")
    tables = caption_numbers(document, "Table")
    require(figures == set(range(1, 17)), f"unexpected figure captions: {sorted(figures)}")
    require(tables == set(range(1, 23)), f"unexpected table captions: {sorted(tables)}")
    svg_paths = sorted(FIGURES.glob("Figure_*.svg"))
    require(len(svg_paths) == 16, f"expected 16 editable SVG figures, found {len(svg_paths)}")
    for path in svg_paths:
        ET.parse(path)
        svg_text = path.read_text(encoding="utf-8", errors="replace").lower()
        require("<image" not in svg_text and "data:image" not in svg_text, f"raster image embedded in SVG: {path.name}")

    main_words = len(re.findall(r"\b[\w-]+\b", body))
    require(main_words <= 10_000, f"main manuscript exceeds JFMI research-paper guidance: {main_words}")

    pdf_info = subprocess.run(["pdfinfo", str(PDF)], check=True, capture_output=True, text=True).stdout
    pdf_text = subprocess.run(["pdftotext", str(PDF), "-"], check=True, capture_output=True, text=True).stdout
    require("JavaScript:      no" in pdf_info, "PDF JavaScript check failed")
    require("Encrypted:       no" in pdf_info, "PDF encryption check failed")
    author_line = next((line for line in pdf_info.splitlines() if line.startswith("Author:")), "")
    require(not author_line.removeprefix("Author:").strip(), "blinded PDF author metadata is not empty")
    for forbidden in (
        "Md Zubaer Mahmood Zubraj",
        "zubraj14@gmail.com",
        "Zubaer Mahmood Zubraj",
        "zmzubraj@gmail.com",
        "0009-0004-3251-3385",
        "Independent Researcher",
    ):
        require(forbidden.lower() not in pdf_text.lower(), f"identity appears in blinded PDF: {forbidden}")

    title_page_text = collect_text(Document(TITLE_PAGE))
    require("Zubaer Mahmood Zubraj" in title_page_text, "confirmed author name missing from title page")
    require("zmzubraj@gmail.com" in title_page_text, "confirmed email missing from title page")
    require("https://orcid.org/0009-0004-3251-3385" in title_page_text, "confirmed ORCID missing from title page")
    require("Self-funded" in title_page_text, "confirmed funding statement missing from title page")
    require("MIT code licence confirmed" in title_page_text, "confirmed MIT code licence missing from title page")
    require("Confirmed by the author on 29 August 2026" in title_page_text, "exclusivity confirmation missing")
    require("THE AUTHOR ANSWERED NO TO APPROVAL" in title_page_text, "AI-use approval refusal is not recorded")
    unresolved_title_page_fields = len(re.findall(r"\[(?:ACCOUNTABLE AUTHOR|AUTHOR|CONFIRM|NOT YET)[^\]]*\]", title_page_text))
    require(unresolved_title_page_fields > 0, "title-page author gate unexpectedly absent")

    cover_letter = Document(COVER_LETTER)
    cover_letter_text = collect_text(cover_letter)
    require("PROVISIONAL OPTIONAL COVER LETTER" in cover_letter_text, "cover-letter gate banner missing")
    require(JFMI_TITLE in cover_letter_text, "publishing title missing from cover letter")
    require("Zubaer Mahmood Zubraj" in cover_letter_text, "confirmed author missing from cover letter")
    require("not been published and is not under consideration" in cover_letter_text, "exclusivity statement missing from cover letter")
    require("Funding is self-funded" in cover_letter_text, "funding statement missing from cover letter")
    require("ACCOUNTABLE AUTHOR MUST INSERT AND APPROVE" in cover_letter_text, "cover-letter declaration gate missing")
    require("systems/cryptography reviewer" not in cover_letter_text.lower(), "deferred reviewer requirement appears in cover letter")
    require("privacy/aml-method reviewer" not in cover_letter_text.lower(), "deferred reviewer requirement appears in cover letter")

    inventory = json.loads(INVENTORY.read_text(encoding="utf-8"))
    require(inventory["submission_ready"] is False, "inventory must remain fail-closed while author gates are open")
    require(inventory["target_status"] == "PROVISIONAL_NOT_CONFIRMED", "inventory incorrectly confirms JFMI")
    require(len(inventory["author_deferred_manual_entry_on_2026_08_29"]) == 5, "author-deferred manual fields are not fully recorded")
    require(len(inventory["open_human_gates"]) >= 10, "inventory omits open human gates")
    for item in inventory["files"]:
        inventory_path = ROOT / item["path"]
        require(inventory_path.exists(), f"inventory path missing: {item['path']}")

    report = {
        "status": "PASS_WITH_AUTHOR_GATES",
        "journal": "Journal of Financial Market Infrastructures",
        "journal_status": "PROVISIONAL_TRIAL_TARGET_NOT_CONFIRMED",
        "official_requirements_checked_on": "2026-08-29",
        "official_requirements_url": "https://www.risk.net/static/risk-journals-submission-guidelines",
        "blinded_docx": str(DOCX.relative_to(ROOT)),
        "blinded_pdf": str(PDF.relative_to(ROOT)),
        "title_page_template": str(TITLE_PAGE.relative_to(ROOT)),
        "optional_cover_letter_template": str(COVER_LETTER.relative_to(ROOT)),
        "submission_inventory": str(INVENTORY.relative_to(ROOT)),
        "abstract_words": abstract_words,
        "keyword_count": keyword_count,
        "key_message_lengths": [len(message) for message in key_messages],
        "author_date_citations": 44,
        "apa_references": len(references),
        "main_text_words_including_table_text": main_words,
        "figures": len(figures),
        "tables": len(tables),
        "editable_svg_figures": len(svg_paths),
        "blinded_identity_check": "PASS",
        "deferred_reviewer_requirement_absent": True,
        "unresolved_title_page_fields": unresolved_title_page_fields,
        "submission_blockers": [
            "final journal confirmation",
            "affiliation and postal address",
            "corresponding-author designation",
            "conflict-of-interest declaration approval",
            "AI-use disclosure approval",
            "public repository URL and DOI/Zenodo deposit",
            "manuscript and figure content-licence decision",
            "accountable human portal review and submission approval",
        ],
    }
    REPORT.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
