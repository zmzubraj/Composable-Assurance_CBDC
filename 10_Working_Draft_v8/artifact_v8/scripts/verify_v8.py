from __future__ import annotations

import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
OUTPUT = ROOT / "output"
DOCX = OUTPUT / "Composable_Assurance_CBDC_Author_Draft_v8.docx"
PDF = OUTPUT / "Composable_Assurance_CBDC_Author_Draft_v8.pdf"
PUBLISHING_TITLE = "Composable assurance for sovereign digital currency (CBDC): An evidence-gated qualification framework"


def require(condition: bool, message: str) -> None:
    if not condition:
        raise RuntimeError(message)


def load_json(name: str) -> dict:
    return json.loads((RESULTS / name).read_text())


def collect_text(document: Document) -> str:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    return text


def cited_reference_numbers(body: str) -> set[int]:
    cited: set[int] = set()
    for start, end in re.findall(r"\[(\d+)\]\s*[-–]\s*\[(\d+)\]", body):
        cited.update(range(int(start), int(end) + 1))
    for group in re.findall(r"\[([0-9,\-\s]+)\]", body):
        for part in group.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                cited.update(range(int(start), int(end) + 1))
            elif part.isdigit():
                cited.add(int(part))
    return cited


def main() -> None:
    require(DOCX.exists() and DOCX.stat().st_size > 100_000, "v8 DOCX missing or unexpectedly small")
    require(PDF.exists() and PDF.stat().st_size > 100_000, "v8 PDF missing or unexpectedly small")
    require((ROOT / "LICENSE-CODE-MIT.txt").is_file(), "MIT code licence missing")
    require((ROOT / "docs" / "FORMAL_MODEL_TO_PROTOTYPE_MAPPING_V8.md").is_file(), "model-to-prototype mapping missing")
    require((ROOT / "docs" / "ANALYSIS_ROLE_AND_DEVIATIONS_V8.md").is_file(), "analysis-role/deviation ledger missing")
    require(not (ROOT / "LICENSE-CODE-APACHE-2.0.txt").exists(), "superseded Apache code licence remains")
    citation_cff = (ROOT / "CITATION.cff").read_text(encoding="utf-8")
    require("license: MIT" in citation_cff, "CITATION.cff does not record MIT")
    require("https://orcid.org/0009-0004-3251-3385" in citation_cff, "CITATION.cff ORCID missing")

    model = load_json("cross_border_model_v5.json")
    require(model["reachable_states"] == 2013, "unexpected reachable-state count")
    require(model["transition_edges"] == 19480, "unexpected transition count")

    privacy = load_json("privacy_v7_experiment_summary.json")
    require(privacy["candidate_sizes"] == [100, 1000, 3000], "privacy candidate sizes changed")

    dp = load_json("dp_v7_summary.json")
    require(dp["l1_sensitivity"] == 3, "DP sensitivity changed")
    require(dp["observed_max_neighbor_l1_over_10000_tests"] == 3, "DP neighbor test changed")
    require(dp["privacy_budget_ledger_required"] is True, "DP budget ledger requirement changed")
    require(dp["annual_release_count"] == 12, "DP annual release count changed")
    require(dp["selected_example_epsilon_total"] == 4, "DP example epsilon changed")

    aml = load_json("aml_v7_summary.json")
    require(aml["independent_graphs"] == 16, "AML graph count changed")
    require(aml["generator_families"] == 4, "AML family count changed")
    require(aml["positive_effect_graphs"] == 12, "AML effect count changed")

    sanctions = load_json("sanctions_v7_summary.json")
    require(sanctions["official_records_indexed"] == 19178, "sanctions record count changed")
    require(sanctions["official_aliases_indexed"] == 20010, "sanctions alias count changed")
    require(sanctions["ownership_monotonicity_properties_passed"] == 500, "ownership property test failed")

    economic = load_json("economic_v7_summary.json")
    require(economic["universal_optimum_identified"] is False, "economic claim boundary changed")

    performance = load_json("performance_v7_summary.json")
    require(performance["national_scale_demonstrated"] is False, "national-scale boundary changed")
    require(performance["scenarios"] == 7, "performance scenario count changed")
    scenarios = pd.read_csv(RESULTS / "performance_v7_scenarios.csv")
    require(int((~scenarios.domestic_slo_pass).sum()) == 1, "expected one overload-control failure")

    prototype = load_json("cross_border_bft_v5.json")
    require(prototype["completed_cross_border_transfers"] == 84, "prototype transfer count changed")
    require(prototype["mismatched_quote_rejected"] is True, "tampered quote was not rejected")
    require(prototype["stale_compliance_rejected"] is True, "stale compliance evidence was not rejected")
    require(prototype["conflicting_abort_quorum_obtained"] is False, "conflicting quorum obtained")
    require(prototype["ledger_A_invariant"]["supply_ok"], "ledger A supply invariant failed")
    require(prototype["ledger_B_invariant"]["supply_ok"], "ledger B supply invariant failed")

    document = Document(DOCX)
    text = collect_text(document)
    normalized_title = " ".join(document.paragraphs[0].text.split())
    require(normalized_title == PUBLISHING_TITLE, "publishing title changed")
    require(document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, "title is not centered")
    require(all(run.bold is True for run in document.paragraphs[0].runs), "title is not bold")
    require(all(run.font.size is not None for run in document.paragraphs[0].runs), "title font size is not explicit")
    require(all(run.font.size.pt == 18 for run in document.paragraphs[0].runs), "title font is not 18 pt")
    require("\n" in document.paragraphs[0].text, "title does not use the intended two-line layout")
    require(document.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER, "author line is not centered")
    require(document.paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.CENTER, "author metadata is not centered")
    require(document.paragraphs[2].runs[0].font.size.pt == 12.5, "author font is not 12.5 pt")
    require(all(run.font.size.pt == 10 for run in document.paragraphs[3].runs), "metadata font is not 10 pt")
    require("\n" in document.paragraphs[3].text, "email and ORCID are not on separate lines")
    require(document.core_properties.author == "Zubaer Mahmood Zubraj", "confirmed author metadata changed")
    require("zmzubraj@gmail.com" in text, "confirmed author email missing")
    require("https://orcid.org/0009-0004-3251-3385" in text, "confirmed ORCID missing")
    require("pending author confirmation" not in text.lower(), "deferred author fields leaked into title-page prose")
    abstract_start = text.index("Abstract") + len("Abstract")
    abstract_end = text.index("Keywords:", abstract_start)
    abstract_words = len(re.findall(r"\b[\w-]+\b", text[abstract_start:abstract_end]))
    require(150 <= abstract_words <= 200, f"abstract word count outside 150-200: {abstract_words}")
    keywords = text[text.index("Keywords:") + len("Keywords:"):].split("\n", 1)[0].split(";")
    require(4 <= len(keywords) <= 6, f"keyword count outside 4-6: {len(keywords)}")

    for forbidden in ["Smart Taka", "ReasonedTransfer", "ERC-20", "ERC-621"]:
        require(forbidden.lower() not in text.lower(), f"forbidden term present: {forbidden}")
    require("systems/cryptography reviewer" not in text.lower(), "deferred reviewer requirement appears in manuscript")
    require("privacy/aml-method reviewer" not in text.lower(), "deferred reviewer requirement appears in manuscript")

    body = text.split("References")[0]
    cited = cited_reference_numbers(body)
    missing = [number for number in range(1, 44) if number not in cited]
    require(not missing, f"uncited references: {missing}")
    for value in ["8.43", "3,571", "1,643"]:
        require(value in text, f"key reported value missing from manuscript: {value}")

    report = {
        "status": "PASS",
        "optimization_safe_verifier": True,
        "docx_bytes": DOCX.stat().st_size,
        "pdf_bytes": PDF.stat().st_size,
        "abstract_words": abstract_words,
        "keyword_count": len(keywords),
        "references_cited": sorted(cited),
        "deferred_reviewer_requirement_absent": True,
        "title_page_format": {
            "centered": True,
            "bold": True,
            "font_pt": 18,
            "title_lines": 2,
            "author_font_pt": 12.5,
            "metadata_font_pt": 10,
            "metadata_lines": 2,
            "deferred_field_notice_absent": True,
        },
        "claim_traceability": {
            "model_to_prototype_mapping": "docs/FORMAL_MODEL_TO_PROTOTYPE_MAPPING_V8.md",
            "analysis_role_and_deviations": "docs/ANALYSIS_ROLE_AND_DEVIATIONS_V8.md",
            "formal_refinement_claimed": False,
        },
        "cross_border_states": model["reachable_states"],
        "prototype_transfers": prototype["completed_cross_border_transfers"],
        "aml_graphs": aml["independent_graphs"],
        "dp_sensitivity": dp["l1_sensitivity"],
        "dp_budget_enforcement_test": "tests/test_privacy_dp.py::test_privacy_accountant_enforces_total_budget",
        "sanctions_records": sanctions["official_records_indexed"],
        "performance_overload_failures": int((~scenarios.domestic_slo_pass).sum()),
    }
    (ROOT / "docs" / "VERIFICATION_V8.json").write_text(json.dumps(report, indent=2) + "\n")
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
